from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document


DEFAULT_REPORT_PATH = Path(r"C:\Users\mekni\Documents\Course\Stage\Memoire\WebSite\Rapport_SmartArchive.docx")

MISSING_CAPTIONS = {
    "Exemple de pièce d’identité utilisée comme entrée documentaire": (
        "Figure 11. Exemple de pièce d’identité utilisée comme entrée documentaire"
    ),
    "Deuxième face ou variante d’entrée documentaire": (
        "Figure 12. Deuxième face ou variante d’entrée documentaire"
    ),
    "Capture d’exemple associée au contexte de traitement administratif": (
        "Figure 13. Capture d’exemple associée au contexte de traitement administratif"
    ),
}


def ensure_backup(report_path: Path) -> Path:
    backup_path = report_path.with_name(f"{report_path.stem}.before_table_des_figures.bak{report_path.suffix}")
    if not backup_path.exists():
        shutil.copy2(report_path, backup_path)
    return backup_path


def normalize_body_captions(report_path: Path) -> None:
    doc = Document(str(report_path))
    changed = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text in MISSING_CAPTIONS:
            paragraph.text = MISSING_CAPTIONS[text]
            paragraph.style = "Caption"
            changed = True
            continue

        if re.match(r"^Figure\s+\d+\.", text):
            if paragraph.style.name != "Caption":
                paragraph.style = "Caption"
                changed = True

    if changed:
        doc.save(str(report_path))


def update_front_figure_list(report_path: Path) -> None:
    ps_script = rf"""
$ErrorActionPreference = 'Stop'
$path = '{report_path}'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open($path)
try {{
  $figurePages = @{{}}
  $frontFigureParagraphs = @()
  $insideFigureList = $false
  $afterAbbreviations = $false

  for ($i = 1; $i -le $doc.Paragraphs.Count; $i++) {{
    $paragraph = $doc.Paragraphs.Item($i)
    $rawText = $paragraph.Range.Text
    $text = ($rawText -replace '[\r\a]', '').Trim()

    if ($text -eq 'Liste des figures') {{
      $insideFigureList = $true
      continue
    }}

    if ($text -eq 'Liste des abréviations') {{
      $insideFigureList = $false
      $afterAbbreviations = $true
      continue
    }}

    if ($insideFigureList -and $text -match '^Figure\s+(\d+)\.') {{
      $frontFigureParagraphs += [PSCustomObject]@{{
        Index = $i
        Text = $text
        Number = [int]$matches[1]
      }}
      continue
    }}

    if ($afterAbbreviations -and $text -match '^Figure\s+(\d+)\.') {{
      $num = [int]$matches[1]
      if (-not $figurePages.ContainsKey($num)) {{
        $figurePages[$num] = $paragraph.Range.Information(3)
      }}
    }}
  }}

  foreach ($item in $frontFigureParagraphs) {{
    if ($figurePages.ContainsKey($item.Number)) {{
      $paragraph = $doc.Paragraphs.Item($item.Index)
      Write-Output ('UPDATING_FIGURE_' + $item.Number + '_PAGE_' + $figurePages[$item.Number])
      $paragraph.Range.Text = $item.Text + "`t" + $figurePages[$item.Number] + "`r"
    }}
  }}

  $doc.Fields.Update() | Out-Null
  $doc.Save()
}} finally {{
  $doc.Close()
  try {{ $word.Quit() }} catch {{}}
}}
exit 0
"""

    completed = subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps_script],
        capture_output=True,
        text=True,
        check=False,
    )

    if completed.returncode != 0:
        raise RuntimeError(
            "Word automation failed while updating the figure table.\n"
            f"STDOUT:\n{completed.stdout}\nSTDERR:\n{completed.stderr}"
        )


def style_front_figure_list(report_path: Path) -> None:
    doc = Document(str(report_path))
    inside_figure_list = False
    changed = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text == "Liste des figures":
            inside_figure_list = True
            continue

        if text == "Liste des abréviations":
            inside_figure_list = False
            continue

        if inside_figure_list and re.match(r"^Figure\s+\d+\.", text):
            if paragraph.style.name != "toc 1":
                paragraph.style = "toc 1"
                changed = True

    if changed:
        doc.save(str(report_path))


def main() -> int:
    report_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_REPORT_PATH

    if not report_path.exists():
        print(f"Missing report: {report_path}")
        return 1

    backup_path = ensure_backup(report_path)
    normalize_body_captions(report_path)
    update_front_figure_list(report_path)
    style_front_figure_list(report_path)

    print(f"Updated report: {report_path}")
    print(f"Backup kept: {backup_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
