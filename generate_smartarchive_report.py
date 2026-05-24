"""
Generate a polished Word report for the SmartArchive project.

Output:
    Rapport_SmartArchive_Final.docx
"""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


PROJECT_DIR = Path(__file__).resolve().parent
BASE_DIR = PROJECT_DIR.parent
OUTPUT = BASE_DIR / "Rapport_SmartArchive_Final.docx"
LOGO = PROJECT_DIR / "src" / "assets" / "SmartArchiveLogo.png"
ADMIN_CAPTURE = BASE_DIR / "Exemples-CIN" / "ADMIN" / "Capture d'écran 2026-04-22 185440.png"
USER_CIN_1 = BASE_DIR / "Exemples-CIN" / "USER" / "IMG_20240513_123230.jpg"
USER_CIN_2 = BASE_DIR / "Exemples-CIN" / "USER" / "IMG_20240513_123241.jpg"
REPORT_ASSETS_DIR = BASE_DIR / "report_assets"

BLUE = RGBColor(0x18, 0x3B, 0x66)
LIGHT_BLUE = RGBColor(0x2C, 0x6F, 0xBF)
GRAY = RGBColor(0x55, 0x55, 0x55)

ARCHI_DIAGRAM = REPORT_ASSETS_DIR / "diagram_architecture.png"
AUTH_DIAGRAM = REPORT_ASSETS_DIR / "diagram_auth_cin.png"
DOCFLOW_DIAGRAM = REPORT_ASSETS_DIR / "diagram_document_flow.png"
ADMIN_DIAGRAM = REPORT_ASSETS_DIR / "diagram_admin_modules.png"
SCRUM_DIAGRAM = REPORT_ASSETS_DIR / "diagram_scrum_cycle.png"
ROADMAP_DIAGRAM = REPORT_ASSETS_DIR / "diagram_sprint_roadmap.png"
USECASE_DIAGRAM = REPORT_ASSETS_DIR / "diagram_use_case_general.png"
CLASS_DIAGRAM = REPORT_ASSETS_DIR / "diagram_class_core.png"
SEQ_AUTH_DIAGRAM = REPORT_ASSETS_DIR / "diagram_sequence_auth.png"
SEQ_UPLOAD_DIAGRAM = REPORT_ASSETS_DIR / "diagram_sequence_upload.png"


def load_font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_centered(draw, xy, text, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, font=font, fill=fill)


def draw_box(draw, xy1, xy2, text, fill, outline, title_font, body_font):
    draw.rounded_rectangle([xy1, xy2], radius=16, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    total_h = 0
    dims = []
    for line in lines:
        font = title_font if line == lines[0] else body_font
        bbox = draw.textbbox((0, 0), line, font=font)
        h = bbox[3] - bbox[1]
        w = bbox[2] - bbox[0]
        dims.append((line, font, w, h))
        total_h += h + 6
    total_h -= 6
    x_center = (xy1[0] + xy2[0]) / 2
    y = (xy1[1] + xy2[1] - total_h) / 2
    for line, font, w, h in dims:
        draw.text((x_center - w / 2, y), line, font=font, fill=(25, 35, 55))
        y += h + 6


def draw_arrow(draw, start, end, fill=(44, 111, 191)):
    draw.line([start, end], fill=fill, width=4)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    if abs(dx) >= abs(dy):
        if dx >= 0:
            tip = end
            left = (end[0] - 14, end[1] - 8)
            right = (end[0] - 14, end[1] + 8)
        else:
            tip = end
            left = (end[0] + 14, end[1] - 8)
            right = (end[0] + 14, end[1] + 8)
    else:
        if dy >= 0:
            tip = end
            left = (end[0] - 8, end[1] - 14)
            right = (end[0] + 8, end[1] - 14)
        else:
            tip = end
            left = (end[0] - 8, end[1] + 14)
            right = (end[0] + 8, end[1] + 14)
    draw.polygon([tip, left, right], fill=fill)


def draw_dashed_line(draw, start, end, fill=(130, 130, 130), dash=10, gap=8, width=2):
    x1, y1 = start
    x2, y2 = end
    if x1 == x2:
        y = y1
        while y < y2:
            draw.line([(x1, y), (x2, min(y + dash, y2))], fill=fill, width=width)
            y += dash + gap
    elif y1 == y2:
        x = x1
        while x < x2:
            draw.line([(x, y1), (min(x + dash, x2), y2)], fill=fill, width=width)
            x += dash + gap


def draw_actor(draw, center_x, top_y, name, font):
    draw.ellipse((center_x - 18, top_y, center_x + 18, top_y + 36), outline=(40, 40, 40), width=3)
    draw.line((center_x, top_y + 36, center_x, top_y + 105), fill=(40, 40, 40), width=3)
    draw.line((center_x - 35, top_y + 58, center_x + 35, top_y + 58), fill=(40, 40, 40), width=3)
    draw.line((center_x, top_y + 105, center_x - 30, top_y + 150), fill=(40, 40, 40), width=3)
    draw.line((center_x, top_y + 105, center_x + 30, top_y + 150), fill=(40, 40, 40), width=3)
    bbox = draw.textbbox((0, 0), name, font=font)
    draw.text((center_x - (bbox[2] - bbox[0]) / 2, top_y + 162), name, font=font, fill=(25, 35, 55))


def draw_ellipse_usecase(draw, xy1, xy2, text, font):
    draw.ellipse([xy1, xy2], outline=(44, 111, 191), width=3, fill=(247, 251, 255))
    lines = text.split("\n")
    total_h = 0
    dims = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        dims.append((line, w, h))
        total_h += h + 3
    total_h -= 3
    cx = (xy1[0] + xy2[0]) / 2
    y = (xy1[1] + xy2[1] - total_h) / 2
    for line, w, h in dims:
        draw.text((cx - w / 2, y), line, font=font, fill=(20, 30, 45))
        y += h + 3


def make_canvas(title):
    REPORT_ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(38, bold=True)
    body_font = load_font(24, bold=False)
    small_font = load_font(20, bold=False)
    draw_centered(draw, (800, 45), title, title_font, (24, 59, 102))
    return img, draw, title_font, body_font, small_font


def generate_architecture_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Architecture globale de SmartArchive")
    draw_box(draw, (80, 180), (420, 420), "Frontend React\nAuthentification\nDashboard\nUpload\nAdministration", (232, 243, 255), (44, 111, 191), body_font, small_font)
    draw_box(draw, (620, 160), (980, 440), "Backend Express\nAPI REST\nAuth\nDocuments\nAdmin Routes", (237, 247, 237), (80, 140, 90), body_font, small_font)
    draw_box(draw, (1160, 180), (1490, 420), "MongoDB\nUsers\nAdmins\nDocumentRecord\nSettings", (255, 244, 230), (215, 140, 40), body_font, small_font)
    draw_box(draw, (620, 560), (980, 790), "Services Python\nOCR CIN\nExtraction\nClassification", (248, 238, 255), (138, 95, 180), body_font, small_font)
    draw_box(draw, (80, 580), (420, 770), "Service Email\nVérification\nRéinitialisation\nNotifications", (255, 240, 240), (180, 85, 85), body_font, small_font)
    draw_arrow(draw, (420, 300), (620, 300))
    draw_arrow(draw, (980, 300), (1160, 300))
    draw_arrow(draw, (800, 440), (800, 560))
    draw_arrow(draw, (620, 670), (420, 670))
    draw.text((485, 285), "HTTP / JSON", font=small_font, fill=(60, 60, 60))
    draw.text((1030, 285), "Mongoose", font=small_font, fill=(60, 60, 60))
    draw.text((820, 490), "OCR / NLP", font=small_font, fill=(60, 60, 60))
    draw.text((450, 655), "SMTP", font=small_font, fill=(60, 60, 60))
    img.save(ARCHI_DIAGRAM)


def generate_auth_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Flux d'authentification basé sur la CIN")
    steps = [
        ((80, 300), (300, 470), "1. Upload CIN\npar l'utilisateur", (232, 243, 255), (44, 111, 191)),
        ((370, 300), (620, 470), "2. OCR CIN\net extraction", (248, 238, 255), (138, 95, 180)),
        ((690, 300), (940, 470), "3. Vérification\ncompte existant", (237, 247, 237), (80, 140, 90)),
        ((1010, 300), (1260, 470), "4. Envoi code\nemail", (255, 240, 240), (180, 85, 85)),
        ((1330, 300), (1540, 470), "5. Vérification\net accès", (255, 244, 230), (215, 140, 40)),
    ]
    for xy1, xy2, text, fill, outline in steps:
        draw_box(draw, xy1, xy2, text, fill, outline, body_font, small_font)
    for i in range(len(steps) - 1):
        draw_arrow(draw, (steps[i][1][0], 385), (steps[i + 1][0][0], 385))
    draw.text((700, 540), "Si nouvel utilisateur : création contrôlée du compte puis code de vérification", font=small_font, fill=(60, 60, 60))
    img.save(AUTH_DIAGRAM)


def generate_document_flow_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Chaîne de traitement documentaire")
    draw_box(draw, (100, 220), (330, 390), "Upload\nfichier", (232, 243, 255), (44, 111, 191), body_font, small_font)
    draw_box(draw, (420, 220), (650, 390), "OCR\ntexte brut", (248, 238, 255), (138, 95, 180), body_font, small_font)
    draw_box(draw, (740, 220), (970, 390), "Extraction\nentités", (237, 247, 237), (80, 140, 90), body_font, small_font)
    draw_box(draw, (1060, 220), (1290, 390), "Classification\n+ score qualité", (255, 244, 230), (215, 140, 40), body_font, small_font)
    draw_box(draw, (1380, 220), (1530, 390), "Stockage\nMongoDB", (255, 240, 240), (180, 85, 85), body_font, small_font)
    draw_box(draw, (500, 560), (860, 760), "Restitution utilisateur\nHistorique, Ma CIN,\nédition des entités", (238, 248, 242), (58, 122, 91), body_font, small_font)
    draw_box(draw, (980, 560), (1360, 760), "Supervision admin\nReview, statistiques,\nrapports", (242, 243, 255), (93, 102, 181), body_font, small_font)
    draw_arrow(draw, (330, 305), (420, 305))
    draw_arrow(draw, (650, 305), (740, 305))
    draw_arrow(draw, (970, 305), (1060, 305))
    draw_arrow(draw, (1290, 305), (1380, 305))
    draw_arrow(draw, (1450, 390), (1180, 560))
    draw_arrow(draw, (1450, 390), (680, 560))
    img.save(DOCFLOW_DIAGRAM)


def generate_admin_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Modules fonctionnels de l'administration")
    draw_box(draw, (580, 120), (1020, 260), "Administrateur SmartArchive", (232, 243, 255), (44, 111, 191), body_font, small_font)
    modules = [
        ((120, 380), (460, 560), "Gestion\nutilisateurs"),
        ((520, 380), (860, 560), "Gestion\ndocuments"),
        ((920, 380), (1260, 560), "Statistiques\net KPIs"),
        ((1320, 380), (1540, 560), "Rapports\net exports"),
        ((520, 650), (860, 820), "Paramètres\nsystème"),
        ((920, 650), (1260, 820), "Révision\nqualité"),
    ]
    for xy1, xy2, text in modules:
        draw_box(draw, xy1, xy2, text, (245, 247, 250), (110, 125, 140), body_font, small_font)
        draw_arrow(draw, (800, 260), ((xy1[0] + xy2[0]) // 2, xy1[1]))
    img.save(ADMIN_DIAGRAM)


def generate_scrum_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Cycle Scrum adapté au projet SmartArchive")
    draw_box(draw, (80, 330), (300, 500), "Product\nBacklog", (232, 243, 255), (44, 111, 191), body_font, small_font)
    draw_box(draw, (370, 330), (620, 500), "Sprint\nPlanning", (248, 238, 255), (138, 95, 180), body_font, small_font)
    draw_box(draw, (700, 330), (960, 500), "Sprint\nBacklog", (237, 247, 237), (80, 140, 90), body_font, small_font)
    draw_box(draw, (1040, 280), (1320, 550), "Sprint\n2 à 3 semaines\nDév + tests", (255, 244, 230), (215, 140, 40), body_font, small_font)
    draw_box(draw, (1380, 330), (1540, 500), "Incrément", (255, 240, 240), (180, 85, 85), body_font, small_font)
    draw_arrow(draw, (300, 415), (370, 415))
    draw_arrow(draw, (620, 415), (700, 415))
    draw_arrow(draw, (960, 415), (1040, 415))
    draw_arrow(draw, (1320, 415), (1380, 415))
    draw_arrow(draw, (1460, 330), (1200, 170))
    draw_arrow(draw, (1200, 170), (880, 170))
    draw_arrow(draw, (880, 170), (620, 330))
    draw.text((1010, 120), "Review + Rétrospective + ajustements", font=small_font, fill=(60, 60, 60))
    img.save(SCRUM_DIAGRAM)


def generate_roadmap_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Découpage prévisionnel des sprints")
    y_positions = [170, 330, 490, 650]
    labels = [
        ("Sprint 1", "Authentification\nVérification email\nGestion des accès"),
        ("Sprint 2", "Upload documents\nOCR\nClassification\nHistorique"),
        ("Sprint 3", "Espace admin\nReview qualité\nStatistiques"),
        ("Sprint 4", "Rapports\nParamètres\nOptimisations"),
    ]
    durations = ["2 semaines", "3 semaines", "2 semaines", "2 semaines"]
    for idx, y in enumerate(y_positions):
        draw.rounded_rectangle((100, y, 350, y + 95), radius=18, fill=(232, 243, 255), outline=(44, 111, 191), width=3)
        draw_centered(draw, (225, y + 48), labels[idx][0], body_font, (24, 59, 102))
        draw.polygon([(390, y + 10), (1180, y + 10), (1260, y + 47), (1180, y + 85), (390, y + 85)], fill=(245, 247, 250), outline=(110, 125, 140))
        draw.multiline_text((430, y + 18), labels[idx][1], font=small_font, fill=(35, 45, 60), spacing=4)
        draw.text((1320, y + 34), durations[idx], font=small_font, fill=(60, 60, 60))
    img.save(ROADMAP_DIAGRAM)


def generate_usecase_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Diagramme de cas d'utilisation général")
    draw.rectangle((250, 120, 1360, 790), outline=(140, 140, 140), width=3)
    draw.text((700, 125), "SmartArchive", font=body_font, fill=(24, 59, 102))
    draw_actor(draw, 120, 240, "Utilisateur", small_font)
    draw_actor(draw, 1480, 240, "Administrateur", small_font)
    usecases = [
        ((500, 180), (820, 260), "S'authentifier"),
        ((420, 310), (860, 390), "Téléverser un document"),
        ((420, 430), (860, 510), "Consulter l'historique"),
        ((420, 550), (860, 630), "Corriger les\nentités extraites"),
        ((940, 270), (1270, 350), "Gérer les utilisateurs"),
        ((940, 390), (1270, 470), "Réviser un document"),
        ((940, 510), (1270, 590), "Consulter les\nstatistiques"),
        ((940, 630), (1270, 710), "Générer des rapports\net configurer"),
    ]
    for xy1, xy2, text in usecases:
        draw_ellipse_usecase(draw, xy1, xy2, text, small_font)
    actor_user = (155, 315)
    for target in [(500, 220), (420, 350), (420, 470), (420, 590)]:
        draw.line([actor_user, target], fill=(40, 40, 40), width=2)
    actor_admin = (1445, 315)
    for target in [(820, 220), (1270, 310), (1270, 430), (1270, 550), (1270, 670)]:
        draw.line([actor_admin, target], fill=(40, 40, 40), width=2)
    img.save(USECASE_DIAGRAM)


def generate_class_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Diagramme de classes principal")
    classes = [
        ((80, 170), (420, 470), "User\n- email\n- password\n- role\n- cin_number\n- is_verified\n+ login()\n+ updateProfile()"),
        ((520, 170), (940, 470), "DocumentRecord\n- original_filename\n- document_type\n- quality_score\n- extracted_data\n- full_text\n- verification_status\n+ save()\n+ updateEntities()"),
        ((1040, 170), (1500, 470), "AdminSettings\n- siteName\n- sessionTimeout\n- autoVerifyThreshold\n- notifyByEmail\n+ updateSettings()"),
        ((300, 580), (700, 820), "Admin\n- email\n- cin_number\n- role=admin\n+ reviewDocument()\n+ generateReport()"),
        ((860, 580), (1280, 820), "VerificationCode\n- email\n- code\n- purpose\n- expires_at\n+ validate()"),
    ]
    for xy1, xy2, text in classes:
        draw_box(draw, xy1, xy2, text, (249, 250, 252), (90, 105, 130), body_font, small_font)
    draw.line((420, 300, 520, 300), fill=(50, 50, 50), width=3)
    draw.text((442, 270), "1..*", font=small_font, fill=(60, 60, 60))
    draw.text((468, 315), "possède", font=small_font, fill=(60, 60, 60))
    draw.line((620, 470, 520, 580), fill=(50, 50, 50), width=3)
    draw.text((535, 500), "contrôle", font=small_font, fill=(60, 60, 60))
    draw.line((760, 470, 960, 580), fill=(50, 50, 50), width=3)
    draw.text((810, 505), "génère / utilise", font=small_font, fill=(60, 60, 60))
    draw.line((940, 300, 1040, 300), fill=(50, 50, 50), width=3)
    draw.text((965, 270), "lit / règle", font=small_font, fill=(60, 60, 60))
    img.save(CLASS_DIAGRAM)


def generate_sequence_auth_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Diagramme de séquence - Authentification par CIN")
    labels = ["Utilisateur", "Frontend", "API /auth", "MongoDB", "Email Service"]
    x = [140, 450, 760, 1080, 1380]
    for i, label in enumerate(labels):
        draw.text((x[i] - 60, 120), label, font=small_font, fill=(24, 59, 102))
        draw.ellipse((x[i] - 20, 160, x[i] + 20, 200), outline=(24, 59, 102), width=3, fill=(247, 251, 255))
        draw_dashed_line(draw, (x[i], 200), (x[i], 760))
    messages = [
        (140, 450, 240, "1. upload CIN"),
        (450, 760, 300, "2. verify / send-code"),
        (760, 1080, 360, "3. check account"),
        (1080, 760, 420, "4. account status"),
        (760, 1380, 480, "5. send email code"),
        (1380, 760, 540, "6. sent"),
        (450, 760, 620, "7. verify-code"),
        (760, 1080, 680, "8. validate code"),
        (1080, 760, 730, "9. ok"),
        (760, 450, 760, "10. token + user"),
    ]
    for sx, ex, y, text in messages:
        draw_arrow(draw, (sx, y), (ex, y), fill=(44, 111, 191))
        draw.text(((sx + ex) / 2 - 70, y - 22), text, font=small_font, fill=(60, 60, 60))
    img.save(SEQ_AUTH_DIAGRAM)


def generate_sequence_upload_diagram():
    img, draw, title_font, body_font, small_font = make_canvas("Diagramme de séquence - Upload et traitement documentaire")
    labels = ["Utilisateur", "Frontend", "API /documents", "OCR Service", "Extractor", "MongoDB"]
    x = [120, 360, 620, 900, 1160, 1450]
    for i, label in enumerate(labels):
        draw.text((x[i] - 55, 120), label, font=small_font, fill=(24, 59, 102))
        draw.ellipse((x[i] - 20, 160, x[i] + 20, 200), outline=(24, 59, 102), width=3, fill=(247, 251, 255))
        draw_dashed_line(draw, (x[i], 200), (x[i], 760))
    messages = [
        (120, 360, 240, "1. choisir fichier"),
        (360, 620, 300, "2. upload"),
        (620, 900, 360, "3. OCR request"),
        (900, 620, 420, "4. text"),
        (620, 1160, 480, "5. extract entities"),
        (1160, 620, 540, "6. entities"),
        (620, 1450, 600, "7. save record"),
        (1450, 620, 660, "8. id / status"),
        (620, 360, 720, "9. response"),
    ]
    for sx, ex, y, text in messages:
        draw_arrow(draw, (sx, y), (ex, y), fill=(44, 111, 191))
        draw.text(((sx + ex) / 2 - 78, y - 22), text, font=small_font, fill=(60, 60, 60))
    img.save(SEQ_UPLOAD_DIAGRAM)


def generate_diagrams():
    generate_architecture_diagram()
    generate_auth_diagram()
    generate_document_flow_diagram()
    generate_admin_diagram()
    generate_scrum_diagram()
    generate_roadmap_diagram()
    generate_usecase_diagram()
    generate_class_diagram()
    generate_sequence_auth_diagram()
    generate_sequence_upload_diagram()


def set_doc_language(run):
    rpr = run._element.get_or_add_rPr()
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "fr-FR")
    lang.set(qn("w:eastAsia"), "fr-FR")
    lang.set(qn("w:bidi"), "ar-TN")
    rpr.append(lang)


def style_run(run, size=11, bold=False, italic=False, color=None, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    set_doc_language(run)


def add_paragraph(doc, text="", *, size=12, bold=False, italic=False, align=None, color=None, space_after=8, first_line=0.75):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    style_run(run, size=20, bold=True, color=BLUE)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        size = 17
        color = BLUE
    elif level == 2:
        size = 14
        color = LIGHT_BLUE
    else:
        size = 12
        color = BLUE
    run = p.add_run(text)
    style_run(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                style_run(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], "2C6FBF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    style_run(run, size=10.5)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_image(doc, path, caption, width_inches=5.8):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        run = cp.add_run(caption)
        style_run(run, size=10.5, italic=True, color=GRAY)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Table des matières. Dans Word, cliquez sur cette zone puis choisissez Mettre à jour le champ."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    style_run(run, size=10, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def chapter_page(doc, chapter_number, title):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(f"Chapitre {chapter_number}")
    style_run(run, size=22, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(120)
    run2 = p2.add_run(title)
    style_run(run2, size=20, bold=True, color=LIGHT_BLUE)


doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.2)
    add_page_number_footer(section)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)

generate_diagrams()


       
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    p.add_run().add_picture(str(LOGO), width=Inches(1.4))

add_paragraph(doc, "République Tunisienne", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, first_line=None)
add_paragraph(doc, "Ministère de l’Enseignement Supérieur et de la Recherche Scientifique", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)
add_paragraph(doc, "Institut Supérieur des Études Technologiques", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)
doc.add_paragraph()
add_paragraph(doc, "Rapport de Projet de Fin d’Études", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, first_line=None)
add_paragraph(doc, "Conception et développement de la plateforme web intelligente SmartArchive", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=LIGHT_BLUE, first_line=None)
add_paragraph(doc, "Numérisation, OCR, classification et archivage sécurisé de documents administratifs", italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)
doc.add_paragraph()
add_table(
    doc,
    ["Rubrique", "Information"],
    [
        ["Étudiant", "Mekni Amir"],
        ["Spécialité", "Développement des Systèmes d’Information"],
        ["Projet", "SmartArchive"],
        ["Dossier analysé", str(BASE_DIR)],
        ["Frontend", "React 19 + Vite + React Router + Axios + Recharts"],
        ["Backend", "Node.js + Express + MongoDB + JWT + services Python"],
        ["Année universitaire", "2025 - 2026"],
    ],
)
add_paragraph(doc, "Ce rapport a été élaboré à partir de l’analyse complète du code source du projet SmartArchive et de ses composants frontend, backend et OCR.", italic=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, first_line=None)

add_page_break(doc)

                     
add_title(doc, "Dédicace")
add_paragraph(doc, "Je dédie ce travail à mes parents pour leurs sacrifices, leur confiance et leur soutien constant tout au long de mon parcours universitaire.", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)
add_paragraph(doc, "Je le dédie également à ma famille, à mes enseignants et à toutes les personnes qui m’ont encouragé à persévérer jusqu’à l’aboutissement de ce projet.", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)
add_paragraph(doc, "Enfin, je l’adresse à tous ceux qui croient que la technologie peut simplifier les démarches, sécuriser l’information et créer de la valeur réelle.", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=None)

add_page_break(doc)
add_title(doc, "Remerciements")
add_paragraph(doc, "J’adresse mes sincères remerciements à toutes les personnes qui ont contribué à la réalisation de ce projet de fin d’études.", first_line=0.75)
add_paragraph(doc, "Je remercie particulièrement les encadrants académiques et professionnels pour leurs conseils, leur disponibilité et leurs remarques constructives, qui ont permis d’améliorer la qualité technique et méthodologique de ce travail.")
add_paragraph(doc, "Je remercie également les membres de l’environnement de stage et l’ensemble des personnes ayant facilité l’accès aux informations nécessaires à la compréhension du besoin métier et à l’évaluation de la solution développée.")
add_paragraph(doc, "Ma gratitude va enfin à ma famille et à mes proches pour leur soutien moral, leur patience et leur confiance durant toute la période du projet.")

add_page_break(doc)
add_title(doc, "Résumé")
add_paragraph(doc, "Ce rapport présente la conception et la réalisation de SmartArchive, une plateforme web intelligente destinée à la numérisation, l’extraction et l’archivage de documents administratifs. Le système s’appuie sur une interface web moderne en React, un backend Express relié à MongoDB, ainsi que des services Python spécialisés pour l’OCR et l’extraction de données.")
add_paragraph(doc, "L’application couvre un flux complet : chargement des pièces, détection et validation de CIN, création ou vérification de compte, OCR, classification documentaire, stockage des métadonnées, consultation d’historique, édition des entités extraites et supervision administrative.")
add_paragraph(doc, "L’étude du code montre une architecture modulaire, une séparation claire entre les rôles utilisateur et administrateur, l’intégration d’emails de vérification et de réinitialisation, ainsi qu’un socle orienté évolutivité pour des usages réels d’archivage numérique.")
add_paragraph(doc, "Mots-clés : SmartArchive, OCR, React, Express, MongoDB, numérisation, archivage, classification documentaire, CIN, sécurité.")

add_page_break(doc)
add_title(doc, "Abstract")
add_paragraph(doc, "This report presents the design and implementation of SmartArchive, an intelligent web platform dedicated to document digitization, data extraction and secure archiving. The solution combines a modern React frontend, an Express and MongoDB backend, and Python services for OCR and document intelligence.")
add_paragraph(doc, "The application manages the full workflow, from identity card upload and account verification to document history, metadata correction, administrative review and reporting.")
add_paragraph(doc, "The analysis of the source code highlights a modular architecture, role-based access control, email-driven verification flows and a scalable foundation for future industrial deployment.")

add_page_break(doc)
add_title(doc, "Table des matières")
add_toc(doc)

add_page_break(doc)
add_title(doc, "Liste des figures")
for item in [
    "Figure 1. Architecture globale de SmartArchive.",
    "Figure 2. Chaîne de traitement documentaire.",
    "Figure 3. Flux d'authentification basé sur la CIN.",
    "Figure 4. Modules fonctionnels de l'administration.",
    "Figure 5. Cycle Scrum adapté au projet SmartArchive.",
    "Figure 6. Découpage prévisionnel des sprints.",
    "Figure 7. Diagramme de cas d'utilisation général.",
    "Figure 8. Diagramme de classes principal.",
    "Figure 9. Diagramme de séquence - Authentification par CIN.",
    "Figure 10. Diagramme de séquence - Upload et traitement documentaire.",
    "Figure 11. Exemple de pièce d’identité utilisée comme entrée documentaire.",
    "Figure 12. Deuxième face ou variante d’entrée documentaire.",
    "Figure 13. Capture d’exemple associée au contexte de traitement administratif.",
]:
    add_bullet(doc, item)

add_page_break(doc)
add_title(doc, "Liste des abréviations")
for sigle, definition in [
    ("API", "Application Programming Interface"),
    ("JWT", "JSON Web Token"),
    ("OCR", "Optical Character Recognition"),
    ("UI", "User Interface"),
    ("CRUD", "Create, Read, Update, Delete"),
    ("REST", "Representational State Transfer"),
    ("CIN", "Carte d’Identité Nationale"),
    ("KPI", "Key Performance Indicator"),
    ("SMTP", "Simple Mail Transfer Protocol"),
    ("IA", "Intelligence Artificielle"),
]:
    p = doc.add_paragraph()
    run = p.add_run(f"{sigle} : ")
    style_run(run, size=12, bold=True, color=BLUE)
    run2 = p.add_run(definition)
    style_run(run2, size=12)

add_page_break(doc)
add_title(doc, "Introduction générale")
add_paragraph(doc, "La transformation numérique des services administratifs impose aujourd’hui des solutions capables de traiter rapidement des documents variés tout en garantissant fiabilité, traçabilité et sécurité. Dans ce contexte, la gestion manuelle de pièces numérisées demeure coûteuse, lente et source d’erreurs, notamment lorsque les documents doivent être classés, extraits, vérifiés puis restitués à différents profils utilisateurs.")
add_paragraph(doc, "Le projet SmartArchive s’inscrit dans cette problématique. Son objectif est de proposer une plateforme web moderne capable de centraliser la collecte de documents, d’automatiser une partie de leur lecture grâce à l’OCR, de structurer les informations extraites et de mettre à disposition un espace de consultation et d’administration complet.")
add_paragraph(doc, "Ce rapport a été construit à partir de l’exploration détaillée du répertoire du projet, qui regroupe un frontend React, un backend Express, une base MongoDB et plusieurs services Python orientés reconnaissance documentaire. L’approche adoptée consiste à présenter le projet non pas comme une simple maquette, mais comme une solution full-stack cohérente, déjà structurée autour de flux utilisateurs réels.")
add_paragraph(doc, "Le document est organisé autour de l’étude du besoin, de l’architecture générale, de la réalisation des composants frontend et backend, des mécanismes d’authentification et de contrôle, des traitements OCR et enfin des perspectives d’amélioration.")

           
chapter_page(doc, 1, "Cadre du projet et analyse des besoins")
add_heading(doc, "1.1 Présentation de l’organisme d’accueil : CNI", 1)
add_paragraph(doc, "Le Centre National de l’Informatique, couramment désigné par le sigle CNI, constitue un acteur historique de l’écosystème numérique tunisien. Dans l’exemple de mémoire fourni, il est présenté comme un établissement public fondé le 30 décembre 1975, doté d’une autonomie financière et d’une mission centrale d’appui aux systèmes d’information publics.")
add_paragraph(doc, "Le CNI intervient dans les domaines de l’informatique, de l’accompagnement technologique et de la modernisation des services publics. Son positionnement en fait un environnement cohérent pour le développement d’une plateforme comme SmartArchive, orientée vers la numérisation, la structuration de l’information et l’archivage sécurisé.")

add_heading(doc, "1.2 Domaine d’activité", 1)
add_paragraph(doc, "Le domaine d’activité du CNI couvre l’appui technique, la mise en œuvre de solutions informatiques, l’accompagnement des administrations et la valorisation des technologies de l’information. Dans la logique de l’exemple, il s’agit d’un organisme fortement lié à la transformation numérique et à la fiabilisation des processus documentaires.")

add_heading(doc, "1.3 Missions du CNI", 1)
for item in [
    "Assurer l’appui aux structures publiques dans la réalisation et l’exploitation de systèmes d’information.",
    "Participer à la mise en œuvre de projets de transformation numérique et d’e-administration.",
    "Proposer des méthodes, normes et pratiques d’ingénierie adaptées aux besoins des organismes accompagnés.",
    "Contribuer à la formation et à la diffusion des compétences dans le domaine informatique.",
    "Renforcer la qualité, la sécurité et la traçabilité des traitements numériques.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.4 Présentation générale du projet", 1)
add_paragraph(doc, "SmartArchive est une plateforme web de gestion documentaire intelligente. Le projet vise à simplifier le traitement de documents administratifs, en particulier les pièces d’identité, à travers un cycle complet comprenant le chargement des images, l’extraction OCR, la classification, l’archivage et l’exploitation des données extraites.")
add_paragraph(doc, "Le dossier analysé contient deux sous-projets principaux : un frontend nommé Frontend-SmartArchive et un backend nommé Backend-SmartArchive. Le premier assure l’expérience utilisateur et la visualisation des données, tandis que le second orchestre la logique métier, la sécurité, la persistance des données et la communication avec des services d’intelligence documentaire.")

add_heading(doc, "1.5 Problématique", 1)
add_paragraph(doc, "Dans de nombreux contextes administratifs, les documents sont encore gérés de manière semi-manuelle. Même lorsqu’ils sont numérisés, l’extraction d’informations pertinentes, la vérification d’identité, la recherche historique et la supervision restent dispersées entre plusieurs outils. Cette fragmentation augmente les délais de traitement et complique la gouvernance des données.")
add_paragraph(doc, "SmartArchive répond à cette problématique en proposant une chaîne unifiée. La plateforme limite les ressaisies, centralise les métadonnées, améliore l’accès à l’information et introduit une couche de contrôle administratif pour les cas à faible qualité ou nécessitant une validation humaine.")

add_heading(doc, "1.6 Objectifs du projet", 1)
for item in [
    "Permettre le téléversement sécurisé de documents d’identité ou de pièces administratives.",
    "Extraire automatiquement les informations textuelles et structurées via OCR.",
    "Classer les documents et attribuer un score de qualité.",
    "Offrir un espace utilisateur pour consulter l’historique, corriger les entités et gérer son profil.",
    "Mettre en place un espace administrateur pour la supervision, la vérification, les statistiques et les rapports.",
    "Rendre la solution extensible grâce à une architecture modulaire et des services séparés.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.7 Acteurs et rôles", 1)
add_table(
    doc,
    ["Acteur", "Rôle principal", "Capacités majeures"],
    [
        ["Utilisateur", "Consommation et suivi", "Se connecter, téléverser, consulter l’historique, corriger des entités, voir sa CIN, gérer son profil"],
        ["Administrateur", "Supervision et contrôle", "Gérer les utilisateurs, examiner les documents, produire des rapports, modifier les paramètres système"],
        ["Services intelligents", "Traitement automatisé", "OCR, extraction d’entités, classification documentaire, autocomplétion assistée"],
    ],
)

add_heading(doc, "1.8 Besoins fonctionnels", 1)
for item in [
    "Authentification multi-profils avec vérification par email et gestion des sessions.",
    "Téléversement de documents avec validation du type et de la taille des fichiers.",
    "Gestion spécifique de la CIN avec détection du numéro et contrôle d’existence en base.",
    "Historique consultable avec recherche, filtrage, pagination et détail documentaire.",
    "Édition des champs extraits et du texte OCR pour correction manuelle.",
    "Visualisation d’une fiche 'Ma CIN' adaptée à l’utilisateur connecté.",
    "Tableaux de bord administrateur avec indicateurs, listes et exports CSV.",
    "Gestion centralisée des paramètres de sécurité, notifications et automatisation.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.9 Besoins non fonctionnels", 1)
for item in [
    "Sécurité des accès via JWT, middleware de protection et séparation des rôles.",
    "Maintenabilité par découpage en composants React, routes Express et modèles Mongoose.",
    "Extensibilité grâce à l’ajout possible de nouveaux types documentaires ou nouveaux services Python.",
    "Ergonomie via une interface moderne, des tableaux de bord, des formulaires guidés et des retours de progression.",
    "Traçabilité par stockage des métadonnées, de l’état de vérification, du score de qualité et des journaux d’action côté administration.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.10 Méthodologie Scrum", 1)
add_paragraph(doc, "Afin de conduire le projet de manière progressive et maîtrisée, une démarche inspirée de Scrum peut être retenue, conformément à l’esprit de votre exemple. Scrum organise le travail sous forme d’itérations courtes appelées sprints. Chaque sprint possède un objectif clair, un ensemble de tâches priorisées et un livrable fonctionnel à la fin du cycle.")
add_paragraph(doc, "Dans le cas de SmartArchive, cette approche convient particulièrement bien, car le produit se décompose naturellement en blocs métier : authentification, traitement des documents, espace utilisateur, espace administrateur et exploitation des résultats. Elle permet aussi d’introduire les services OCR et les raffinements d’interface de manière incrémentale.")
for item in [
    "Product Backlog : liste ordonnée des besoins et fonctionnalités à implémenter.",
    "Sprint Planning : sélection des éléments à réaliser durant le sprint.",
    "Sprint Backlog : tâches détaillées retenues pour l’itération en cours.",
    "Review et Rétrospective : validation du livrable et amélioration continue.",
]:
    add_bullet(doc, item)
add_image(doc, SCRUM_DIAGRAM, "Figure 5. Cycle Scrum adapté au projet SmartArchive", 6.6)

add_heading(doc, "1.11 Backlog produit", 1)
add_paragraph(doc, "À partir du code source analysé, il est possible de reconstruire un backlog produit cohérent avec les fonctionnalités réellement développées.")
add_table(
    doc,
    ["ID", "User Story", "Priorité", "Complexité"],
    [
        ["US1", "En tant qu’utilisateur, je peux me connecter et vérifier mon identité via ma CIN.", "Haute", "Moyenne"],
        ["US2", "En tant qu’utilisateur, je peux téléverser un document et suivre son traitement.", "Haute", "Moyenne"],
        ["US3", "En tant qu’utilisateur, je peux consulter l’historique et corriger les entités extraites.", "Haute", "Moyenne"],
        ["US4", "En tant qu’administrateur, je peux gérer les utilisateurs et consulter leurs informations.", "Haute", "Moyenne"],
        ["US5", "En tant qu’administrateur, je peux approuver ou rejeter un document faible qualité.", "Haute", "Moyenne"],
        ["US6", "En tant qu’administrateur, je peux consulter des statistiques et générer des rapports.", "Moyenne", "Élevée"],
        ["US7", "En tant qu’administrateur, je peux configurer les paramètres globaux du système.", "Moyenne", "Moyenne"],
        ["US8", "En tant que système, je peux lancer OCR, extraction et classification automatiquement.", "Haute", "Élevée"],
    ],
)

add_heading(doc, "1.12 Découpage des sprints", 1)
add_paragraph(doc, "Le découpage ci-dessous synthétise une organisation possible du projet en sprints cohérente avec la structure observée dans le dépôt.")
add_table(
    doc,
    ["Sprint", "Objectif principal", "Fonctionnalités associées"],
    [
        ["Sprint 1", "Authentification et accès", "CIN upload, vérification email, login, reset password, profil courant"],
        ["Sprint 2", "Traitement documentaire utilisateur", "Upload documents, OCR, extraction, historique, Ma CIN, correction entités"],
        ["Sprint 3", "Administration", "Dashboard admin, gestion utilisateurs, review documents, statistiques"],
        ["Sprint 4", "Pilotage et finalisation", "Rapports, exports, paramètres système, autocomplétion, optimisations"],
    ],
)
add_image(doc, ROADMAP_DIAGRAM, "Figure 6. Découpage prévisionnel des sprints", 6.6)

add_heading(doc, "1.13 Frameworks, technologies et outils", 1)
add_table(
    doc,
    ["Catégorie", "Outils / technologies", "Usage dans le projet"],
    [
        ["Framework frontend", "React 19 + Vite", "Construction de l’interface web et bundling moderne"],
        ["Navigation", "React Router DOM", "Protection et organisation des routes utilisateur et admin"],
        ["Appels API", "Axios", "Communication HTTP entre frontend et backend"],
        ["Visualisation", "Recharts", "Statistiques et graphiques du module admin"],
        ["Framework backend", "Node.js + Express 5", "API REST, logique métier et middlewares"],
        ["Base de données", "MongoDB + Mongoose", "Persistance structurée des documents, comptes et réglages"],
        ["Sécurité", "JWT, Helmet, bcryptjs / CryptoJS", "Authentification, headers sécurisés, gestion des mots de passe"],
        ["Upload", "Multer, Sharp", "Téléversement et traitement de fichiers"],
        ["Emails", "Nodemailer", "Vérification email, reset password, notifications"],
        ["Services IA/OCR", "Python, Flask, LLMWhisperer", "Reconnaissance de texte et extraction intelligente"],
        ["Outils d’export", "jsPDF, html2canvas, PapaParse", "Export PDF, captures et CSV"],
        ["UI / style", "Lucide React, Headless UI, DotLottie", "Icônes, composants interactifs et animations"],
    ],
)
add_image(doc, ARCHI_DIAGRAM, "Figure 1. Architecture globale de SmartArchive", 6.6)

           
chapter_page(doc, 2, "Architecture générale de SmartArchive")
add_heading(doc, "2.1 Organisation du dépôt", 1)
add_paragraph(doc, "L’arborescence principale étudiée contient un frontend, un backend, un dossier d’exemples de CIN ainsi que des scripts de génération documentaire. Le frontend concentre les pages et composants React. Le backend regroupe les routes, modèles, services et scripts Python nécessaires aux traitements intelligents.")
add_table(
    doc,
    ["Répertoire", "Contenu", "Rôle dans la solution"],
    [
        ["Frontend-SmartArchive", "src, assets, composants, routes", "Couche de présentation et interactions utilisateur"],
        ["Backend-SmartArchive", "routes, models, services, server.js", "Couche métier, sécurité, accès données"],
        ["Exemples-CIN", "captures et exemples d’images", "Jeux d’exemple pour démonstration documentaire"],
    ],
)

add_heading(doc, "2.2 Architecture logique", 1)
add_paragraph(doc, "L’architecture suit une logique client-serveur enrichie par des microservices Python. Le navigateur consomme l’application React, qui dialogue avec le backend Express via des requêtes HTTP. Le backend centralise l’authentification, la validation des droits, l’enregistrement des documents et l’appel aux services OCR ou d’extraction. MongoDB persiste les comptes, les documents, les paramètres et les codes temporaires.")
add_paragraph(doc, "Cette séparation présente plusieurs avantages : indépendance relative du frontend, scalabilité potentielle du backend, spécialisation des services Python et meilleure lisibilité du code par domaine fonctionnel.")
add_image(doc, DOCFLOW_DIAGRAM, "Figure 2. Chaîne de traitement documentaire", 6.6)

add_heading(doc, "2.3 Diagramme de cas d’utilisation général", 1)
add_paragraph(doc, "Le diagramme de cas d’utilisation suivant synthétise les interactions principales entre les acteurs et la plateforme. Il met en évidence la dualité entre le parcours métier de l’utilisateur et les fonctions de supervision réservées à l’administrateur.")
add_image(doc, USECASE_DIAGRAM, "Figure 7. Diagramme de cas d'utilisation général", 6.6)

add_heading(doc, "2.4 Diagramme de classes principal", 1)
add_paragraph(doc, "Le diagramme de classes ci-dessous reconstruit les entités les plus structurantes du backend à partir des modèles Mongoose observés dans le code. Il présente les classes métier dominantes et leurs relations principales.")
add_image(doc, CLASS_DIAGRAM, "Figure 8. Diagramme de classes principal", 6.6)

add_heading(doc, "2.5 Diagrammes de séquences détaillés", 1)
add_paragraph(doc, "Afin de rapprocher davantage ce rapport de l’esprit du mémoire exemple, deux séquences représentatives ont été ajoutées : l’authentification centrée sur la CIN et le scénario de téléversement avec OCR et extraction.")
add_image(doc, SEQ_AUTH_DIAGRAM, "Figure 9. Diagramme de séquence - Authentification par CIN", 6.6)
add_image(doc, SEQ_UPLOAD_DIAGRAM, "Figure 10. Diagramme de séquence - Upload et traitement documentaire", 6.6)

add_heading(doc, "2.6 Flux général de fonctionnement", 1)
for item in [
    "L’utilisateur charge une image de CIN ou un autre document depuis l’interface web.",
    "Le frontend envoie les fichiers aux routes dédiées du backend via Axios.",
    "Le backend stocke temporairement le fichier, appelle le service OCR, récupère un texte brut puis des données structurées.",
    "Les résultats sont enrichis par une classification, un score de qualité et des règles de validation.",
    "Les informations sont enregistrées dans MongoDB puis restituées dans l’espace utilisateur ou administrateur.",
]:
    add_numbered(doc, item)

add_heading(doc, "2.7 Routage principal côté frontend", 1)
add_table(
    doc,
    ["Type", "Routes majeures", "Description"],
    [
        ["Publiques", "/, /login, /register, /verify-email, /forgot-password", "Entrée dans la plateforme et récupération d’accès"],
        ["Utilisateur", "/dashboard, /upload, /documents, /my-cin, /notifications, /settings, /profile", "Parcours métier de l’utilisateur connecté"],
        ["Administrateur", "/admin/dashboard, /admin/users, /admin/documents, /admin/statistics, /admin/reports, /admin/settings", "Pilotage et supervision globale"],
    ],
)

add_heading(doc, "2.8 API principales côté backend", 1)
add_table(
    doc,
    ["Namespace", "Exemples d’endpoints", "Usage"],
    [
        ["/api/auth", "login, me, send-code, verify-code, verify-email, reset-password", "Authentification et gestion des identités"],
        ["/api/cin", "upload, check-cin, records", "Traitement spécialisé de la carte d’identité"],
        ["/api/documents", "upload, history, stats, my-cin, :id, :id/entities", "Cycle de vie documentaire"],
        ["/api/admin", "users, statistics, reports, settings, documents/:id/review", "Administration de la plateforme"],
    ],
)

add_heading(doc, "2.9 Modèle de données", 1)
add_paragraph(doc, "Les principaux modèles Mongoose identifiés sont User, Admin, DocumentRecord et AdminSettings. Le modèle User stocke les informations d’identité, de profil et de rôle. DocumentRecord constitue le noyau documentaire, avec les métadonnées du fichier, le type, les données extraites, le texte OCR, le score de qualité et l’état de révision. AdminSettings centralise les préférences globales comme la durée de session, l’automatisation ou les notifications.")

add_heading(doc, "2.10 Sécurité et contrôle d’accès", 1)
for item in [
    "Protection des routes sensibles par JWT.",
    "Chargement du profil courant via /auth/me.",
    "Redirection frontend selon le rôle pour les routes protégées.",
    "Middleware administrateur pour les opérations critiques.",
    "Gestion de codes de vérification envoyés par email.",
]:
    add_bullet(doc, item)
add_paragraph(doc, "L’étude montre néanmoins un point de vigilance important : certains mots de passe utilisateurs sont chiffrés de manière réversible via AES au lieu d’être exclusivement hachés. Ce choix peut être conservé comme transition fonctionnelle, mais il devra être remplacé par un hachage fort irréversible pour renforcer la sécurité globale.")

           
chapter_page(doc, 3, "Réalisation du frontend")
add_heading(doc, "3.1 Pile technique frontend", 1)
add_paragraph(doc, "Le frontend est construit avec React 19 et Vite. Cette combinaison apporte une expérience de développement fluide, un rendu moderne et une structuration claire autour de composants réutilisables. Le routage s’appuie sur React Router DOM et les appels réseau sur Axios.")
add_heading(doc, "3.2 Authentification et expérience d’entrée", 1)
add_paragraph(doc, "Le composant principal d’authentification adopte un parcours original centré sur la CIN. L’utilisateur peut déposer jusqu’à deux images. Le système tente d’abord d’identifier la personne à partir de la carte d’identité, puis déclenche l’envoi d’un code email pour finaliser l’accès. Cette approche diffère d’un simple formulaire classique et renforce l’identité métier de SmartArchive.")
for item in [
    "Drag and drop d’images avec validation côté interface.",
    "Étapes conditionnelles selon le statut utilisateur ou administrateur.",
    "Gestion d’erreurs et de dialogues de confirmation.",
    "Navigation automatique vers l’espace adéquat après validation.",
]:
    add_bullet(doc, item)
add_image(doc, AUTH_DIAGRAM, "Figure 3. Flux d'authentification basé sur la CIN", 6.6)

add_heading(doc, "3.3 Tableau de bord utilisateur", 1)
add_paragraph(doc, "Le tableau de bord utilisateur agrège des indicateurs de volume, un historique récent et des actions rapides. Il s’appuie sur les endpoints /documents/history et /documents/stats et offre une vue synthétique sur les interactions de l’utilisateur avec la plateforme.")

add_heading(doc, "3.4 Téléversement et historique documentaire", 1)
add_paragraph(doc, "Le composant UploadDocument permet de charger plusieurs catégories de fichiers, notamment CIN, passeport, permis ou certificat. Il affiche la progression, contrôle le format et la taille des pièces, puis envoie les fichiers un par un au backend.")
add_paragraph(doc, "La page d’historique enrichit cette fonctionnalité grâce à un moteur de recherche, des filtres, une pagination, un aperçu du contenu OCR et une possibilité d’édition des données extraites. Cette dernière capacité est particulièrement intéressante car elle ouvre la voie à une logique de post-correction humaine après automatisation.")

add_heading(doc, "3.5 Espace personnel", 1)
for item in [
    "Page Profile pour mettre à jour les informations personnelles.",
    "Page Settings pour modifier le mot de passe et les préférences locales.",
    "Page Notifications basée sur l’historique et localStorage.",
    "Page MyCIN pour afficher une carte synthétique et imprimable issue des données extraites.",
]:
    add_bullet(doc, item)

add_heading(doc, "3.6 Interface d’administration", 1)
add_paragraph(doc, "Le frontend d’administration constitue une véritable application dans l’application. On y retrouve un tableau de bord, un module de gestion des utilisateurs, une gestion des documents, une page de détail, des rapports analytiques et des paramètres système. L’usage de composants réutilisables comme ProTable et FilterPanel améliore la cohérence et la maintenabilité.")
add_paragraph(doc, "Le module Reports mérite une attention particulière : il combine visualisation via Recharts, filtres avancés et assistance à la rédaction avec autocomplétion, ce qui donne au projet une dimension d’aide à l’exploitation et non seulement d’archivage.")
add_image(doc, ADMIN_DIAGRAM, "Figure 4. Modules fonctionnels de l'administration", 6.6)

add_heading(doc, "3.7 Qualités observées dans le frontend", 1)
for item in [
    "Découpage fonctionnel clair entre authentification, utilisateur et administration.",
    "Richesse ergonomique et feedback visuel sur les étapes de traitement.",
    "Bonne exploitation d’outils modernes pour les tableaux, exports et graphiques.",
    "Cohérence générale du parcours utilisateur autour de la CIN et du document.",
]:
    add_bullet(doc, item)

           
chapter_page(doc, 4, "Réalisation du backend et services métier")
add_heading(doc, "4.1 Serveur Express et configuration", 1)
add_paragraph(doc, "Le serveur principal repose sur Express 5. Il charge les variables d’environnement, active CORS, parse les corps JSON, expose le répertoire des fichiers chargés et publie un endpoint /health. Les routes sont découpées par domaine : authentification, upload documentaire, traitement CIN et administration.")

add_heading(doc, "4.2 Authentification et identité", 1)
add_paragraph(doc, "Les routes d’authentification assurent plusieurs cas d’usage : connexion, chargement du profil courant, envoi de code de vérification, validation du code, vérification d’email, réinitialisation de mot de passe et mise à jour du profil. Le service email, basé sur Nodemailer, permet d’envoyer aussi bien des messages de vérification que des notifications de bienvenue.")
add_paragraph(doc, "Le flux le plus intéressant est celui des nouveaux utilisateurs détectés à partir de la CIN. Lorsqu’une carte est reconnue et qu’aucun compte n’existe encore, la plateforme peut créer un compte avec mot de passe temporaire et demander une vérification email. Cette logique rapproche l’inscription d’un enrôlement intelligent plutôt que d’une simple création manuelle.")

add_heading(doc, "4.3 Gestion documentaire", 1)
add_paragraph(doc, "La route /documents/upload constitue le cœur opérationnel du backend. Après réception du fichier, le système déclenche une chaîne de traitement complète : stockage, appel OCR, appel extracteur, classification, détection de langue, évaluation de qualité, puis persistance dans MongoDB.")
for item in [
    "Prise en charge des fichiers via Multer.",
    "Fallback mock si les services OCR ou extracteur sont indisponibles.",
    "Détermination d’un type documentaire et d’un score de qualité.",
    "Marquage requires_admin_review pour les documents faibles ou ambigus.",
    "Exposition d’API de consultation, modification et suppression selon le rôle.",
]:
    add_bullet(doc, item)

add_heading(doc, "4.4 Gestion spécialisée de la CIN", 1)
add_paragraph(doc, "Le namespace /api/cin introduit une logique dédiée à la carte d’identité nationale. Le backend y valide le format du numéro de CIN, détecte l’existence éventuelle d’un utilisateur ou d’un administrateur, alimente ou met à jour un enregistrement documentaire et renvoie des informations directement exploitables par le frontend d’authentification.")
add_paragraph(doc, "Cette spécialisation métier est une force du projet, car elle montre que SmartArchive ne traite pas les documents comme de simples fichiers, mais comme des pièces porteuses d’identité et de règles de contrôle.")

add_heading(doc, "4.5 Administration, reporting et paramétrage", 1)
add_paragraph(doc, "Les routes d’administration permettent de lister et modifier les utilisateurs, examiner des documents, consulter des statistiques, produire des rapports et gérer des paramètres système persistants. Une partie des rapports intègre des suggestions intelligentes et une autocomplétion basée sur un script Python, avec fallback JavaScript si l’environnement externe n’est pas disponible.")

add_heading(doc, "4.6 Modèles de données backend", 1)
add_table(
    doc,
    ["Modèle", "Champs importants", "Apport métier"],
    [
        ["User", "email, password, role, cin_number, is_verified", "Gestion des comptes utilisateur"],
        ["Admin", "email, cin_number, is_verified, role", "Gestion des comptes administrateur"],
        ["DocumentRecord", "document_type, quality_score, extracted_data, full_text, verification_status", "Traçabilité documentaire complète"],
        ["AdminSettings", "sessionTimeout, enforce2FA, autoVerifyThreshold, locale", "Pilotage global de la plateforme"],
    ],
)

           
chapter_page(doc, 5, "Services OCR, extraction et intelligence documentaire")
add_heading(doc, "5.1 Service OCR Python", 1)
add_paragraph(doc, "Le service OCR est implémenté en Python avec Flask et CORS. Il s’appuie sur LLMWhisperer et comporte des traitements spécifiques au texte arabe et aux cartes d’identité. La présence d’un nettoyage lexical dédié montre un effort pour adapter la chaîne de reconnaissance à des documents réels et non à des cas purement académiques.")
add_paragraph(doc, "Le service expose un endpoint de traitement et peut vérifier la disponibilité d’un moteur externe comme Ollama. Cette architecture autorise à terme le remplacement du moteur OCR ou l’ajout de nouveaux modèles sans refondre le backend Node.js.")

add_heading(doc, "5.2 Service de classification et extraction", 1)
add_paragraph(doc, "Un second service Python se charge de la classification et de l’extraction structurée. Il reconnaît plusieurs catégories, notamment la CIN, le passeport, le permis de conduire et d’autres documents. Le backend utilise ensuite cette information pour enrichir le dossier documentaire et faciliter les contrôles humains.")

add_heading(doc, "5.3 Stratégie hybride et tolérance aux pannes", 1)
add_paragraph(doc, "L’une des qualités notables du projet est l’existence de mécanismes de repli. Si un service OCR ou extracteur devient indisponible, le backend peut utiliser des réponses mock. Ce comportement n’est pas une solution finale de production, mais il est très utile en phase de développement et de démonstration, car il évite de bloquer l’ensemble de la plateforme.")

add_heading(doc, "5.4 Apport de l’intelligence dans SmartArchive", 1)
for item in [
    "Détection et validation automatique du numéro de CIN.",
    "Extraction d’informations textuelles à partir d’images.",
    "Classification documentaire semi-automatique.",
    "Aide à la rédaction et autocomplétion pour certains rapports administratifs.",
    "Préparation d’un contrôle humain plutôt qu’automatisation aveugle.",
]:
    add_bullet(doc, item)

add_image(doc, USER_CIN_1, "Exemple de pièce d’identité utilisée comme entrée documentaire", 4.8)
add_image(doc, USER_CIN_2, "Deuxième face ou variante d’entrée documentaire", 4.8)

           
chapter_page(doc, 6, "Tests, résultats et évaluation")
add_heading(doc, "6.1 Scénarios validés par l’étude du code", 1)
add_table(
    doc,
    ["ID", "Scénario", "Résultat observé"],
    [
        ["T1", "Détection d’une CIN existante", "Le frontend adapte le parcours et déclenche un envoi de code"],
        ["T2", "Création d’un nouvel utilisateur à partir d’une CIN reconnue", "Le backend prépare un compte et une vérification email"],
        ["T3", "Upload d’un document administratif", "Le backend exécute OCR, extraction, classification et persistance"],
        ["T4", "Consultation de l’historique", "L’utilisateur visualise ses documents avec filtres et pagination"],
        ["T5", "Révision d’un document faible qualité", "L’administrateur peut approuver ou rejeter le document"],
        ["T6", "Mise à jour des paramètres système", "Les réglages sont persistés via AdminSettings"],
    ],
)

add_heading(doc, "6.2 Résultats fonctionnels obtenus", 1)
add_paragraph(doc, "Au terme de l’analyse, SmartArchive apparaît comme une application fonctionnellement aboutie sur plusieurs dimensions : authentification avancée, gestion multi-rôles, traitement intelligent des pièces, historique détaillé, supervision administrative et outils de pilotage. Le projet n’est pas limité à une interface démonstrative ; il porte déjà une logique métier solide.")

add_heading(doc, "6.3 Indicateurs de volumétrie", 1)
add_table(
    doc,
    ["Indicateur", "Valeur observée"],
    [
        ["Fichiers frontend pertinents", "95"],
        ["Fichiers backend pertinents", "22"],
        ["Volume frontend estimé", "12 282 lignes"],
        ["Volume backend estimé", "9 816 lignes"],
        ["Composants/pages React repérés", "plus de 50"],
        ["Services Python identifiés", "2 principaux + scripts auxiliaires"],
    ],
)

add_heading(doc, "6.4 Illustrations disponibles", 1)
add_image(doc, ADMIN_CAPTURE, "Capture d’exemple associée au contexte de traitement administratif", 6.2)

add_heading(doc, "6.5 Analyse critique", 1)
for item in [
    "Le projet est riche fonctionnellement et montre une réelle cohérence entre frontend et backend.",
    "L’orientation métier autour de la CIN rend la solution originale et pertinente.",
    "Le recours à des fallbacks mock facilite le développement mais devra être encadré en production.",
    "La sécurité des mots de passe doit évoluer vers un hachage irréversible généralisé.",
    "Certaines chaînes de branding ou messages techniques gagneraient à être harmonisées.",
]:
    add_bullet(doc, item)

           
chapter_page(doc, 7, "Perspectives d’amélioration")
add_heading(doc, "7.1 Améliorations de sécurité", 1)
for item in [
    "Remplacer tout chiffrement réversible de mot de passe par un hachage fort type bcrypt.",
    "Mettre en place une authentification multifacteur côté administration.",
    "Tracer plus finement les opérations sensibles avec audit horodaté.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.2 Améliorations techniques", 1)
for item in [
    "Conteneuriser les services Node.js et Python avec Docker.",
    "Ajouter une chaîne CI/CD pour les tests et déploiements.",
    "Introduire une file de messages pour les traitements OCR lourds.",
    "Prévoir un stockage objet externe pour les fichiers volumineux.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.3 Améliorations fonctionnelles", 1)
for item in [
    "Étendre la bibliothèque de types documentaires reconnus.",
    "Ajouter une recherche plein texte sur le contenu OCR.",
    "Mettre en place un tableau de bord d’erreurs de traitement.",
    "Prévoir une version mobile ou responsive plus spécialisée pour la capture terrain.",
]:
    add_bullet(doc, item)

add_heading(doc, "7.4 Perspective métier", 1)
add_paragraph(doc, "À plus long terme, SmartArchive pourrait devenir un noyau de plateforme documentaire pour plusieurs organisations. Grâce à son architecture modulaire, la solution pourrait être étendue à des scénarios de conformité, de vérification inter-services, de certification ou de rapprochement automatique entre documents hétérogènes.")

add_page_break(doc)
add_title(doc, "Conclusion générale")
add_paragraph(doc, "L’étude détaillée du projet SmartArchive montre la réalisation d’une plateforme web sérieuse, moderne et techniquement cohérente. Le système ne se contente pas d’archiver des documents ; il structure un vrai parcours intelligent autour de la capture, de la vérification d’identité, de l’extraction de données et du contrôle administratif.")
add_paragraph(doc, "La richesse du frontend React, la densité fonctionnelle du backend Express, l’utilisation de MongoDB pour la persistance et l’intégration de services Python pour l’OCR donnent au projet une dimension full-stack complète. Cette complémentarité constitue l’un de ses points forts majeurs.")
add_paragraph(doc, "Au-delà de ses qualités actuelles, SmartArchive offre une base robuste pour des évolutions futures. En renforçant certains aspects de sécurité et d’industrialisation, le projet peut être transformé en une solution d’archivage numérique à forte valeur ajoutée pour un usage institutionnel ou professionnel.")

add_page_break(doc)
add_title(doc, "Bibliographie et webographie")
add_heading(doc, "Références institutionnelles et métier", 1)
for ref in [
    "Centre National de l’Informatique (CNI) – https://www.cni.tn/",
]:
    add_bullet(doc, ref)

add_heading(doc, "Références techniques", 1)
for ref in [
    "React Documentation – https://react.dev/",
    "Vite Documentation – https://vite.dev/",
    "React Router Documentation – https://reactrouter.com/",
    "Axios Documentation – https://axios-http.com/",
    "Express Documentation – https://expressjs.com/",
    "Mongoose Documentation – https://mongoosejs.com/",
    "MongoDB Documentation – https://www.mongodb.com/docs/",
    "Nodemailer Documentation – https://nodemailer.com/",
    "Flask Documentation – https://flask.palletsprojects.com/",
    "JSON Web Tokens – https://jwt.io/",
    "Recharts Documentation – https://recharts.org/",
    "jsPDF – https://github.com/parallax/jsPDF",
    "PapaParse – https://www.papaparse.com/",
]:
    add_bullet(doc, ref)

add_heading(doc, "Références de design et style utilisées", 1)
for ref in [
    "Lucide React – https://lucide.dev/",
    "Headless UI React – https://headlessui.com/react",
    "LottieFiles DotLottie React – https://developers.lottiefiles.com/docs/dotlottie-player/dotlottie-react/",
    "HTML2Canvas – https://html2canvas.hertzen.com/",
]:
    add_bullet(doc, ref)

add_page_break(doc)
add_title(doc, "Annexe A - Synthèse technique")
add_table(
    doc,
    ["Élément", "Synthèse"],
    [
        ["Authentification", "JWT, vérification email, réinitialisation, profil courant"],
        ["Documents", "Upload, OCR, classification, historique, édition d’entités"],
        ["CIN", "Analyse spécialisée, validation du numéro, couplage avec le compte"],
        ["Administration", "Utilisateurs, documents, statistiques, rapports, paramètres"],
        ["Services externes", "OCR Python, extracteur, autocomplétion IA"],
    ],
)

final_output = OUTPUT
saved = False
for candidate in [OUTPUT] + [OUTPUT.with_name(f"Rapport_SmartArchive_Final_v{i}.docx") for i in range(2, 10)]:
    try:
        doc.save(str(candidate))
        final_output = candidate
        saved = True
        break
    except PermissionError:
        continue
if not saved:
    raise PermissionError("Impossible de sauvegarder le rapport : tous les fichiers cibles semblent verrouillés.")
print(f"Rapport généré : {final_output}")
