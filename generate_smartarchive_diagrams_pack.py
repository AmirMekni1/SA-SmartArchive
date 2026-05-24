"""
Generate a complete SmartArchive diagrams pack.

Outputs:
    - report_assets/full_diagrams/*.png
    - report_assets/full_diagrams/DIAGRAMS_INDEX.md
    - report_assets/full_diagrams/diagrams_manifest.json
    - Annexe_Diagrammes_SmartArchive.docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


PROJECT_DIR = Path(__file__).resolve().parent
BASE_DIR = PROJECT_DIR.parent
ASSET_DIR = BASE_DIR / "report_assets" / "full_diagrams"
DOCX_OUTPUT = BASE_DIR / "Annexe_Diagrammes_SmartArchive.docx"
INDEX_OUTPUT = ASSET_DIR / "DIAGRAMS_INDEX.md"
MANIFEST_OUTPUT = ASSET_DIR / "diagrams_manifest.json"

PAGE_BG = (255, 255, 255)
TITLE = (22, 56, 96)
TEXT = (35, 42, 52)
MUTED = (100, 110, 125)
BLUE = (46, 111, 191)
BLUE_FILL = (233, 243, 255)
GREEN = (74, 133, 104)
GREEN_FILL = (235, 247, 240)
AMBER = (203, 142, 48)
AMBER_FILL = (255, 244, 225)
ROSE = (185, 90, 98)
ROSE_FILL = (255, 239, 241)
VIOLET = (126, 96, 179)
VIOLET_FILL = (244, 239, 255)
SLATE = (102, 112, 122)
SLATE_FILL = (243, 246, 249)
NOTE_FILL = (255, 248, 205)
LINE = (110, 125, 145)


def load_font(size: int, bold: bool = False, italic: bool = False):
    candidates = []
    if bold and italic:
        candidates.extend(
            [
                "C:/Windows/Fonts/calibriz.ttf",
                "C:/Windows/Fonts/arialbi.ttf",
            ]
        )
    elif bold:
        candidates.extend(
            [
                "C:/Windows/Fonts/calibrib.ttf",
                "C:/Windows/Fonts/arialbd.ttf",
                "C:/Windows/Fonts/timesbd.ttf",
            ]
        )
    elif italic:
        candidates.extend(
            [
                "C:/Windows/Fonts/calibrii.ttf",
                "C:/Windows/Fonts/ariali.ttf",
                "C:/Windows/Fonts/timesi.ttf",
            ]
        )
    else:
        candidates.extend(
            [
                "C:/Windows/Fonts/calibri.ttf",
                "C:/Windows/Fonts/arial.ttf",
                "C:/Windows/Fonts/times.ttf",
            ]
        )

    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = load_font(38, bold=True)
FONT_SUBTITLE = load_font(22, italic=True)
FONT_BODY = load_font(22)
FONT_SMALL = load_font(18)
FONT_TINY = load_font(16)
FONT_BOLD = load_font(22, bold=True)
FONT_SMALL_BOLD = load_font(18, bold=True)


def ensure_dir():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def choose_output_path(target: Path) -> Path:
    if not target.exists():
        return target
    stem = target.stem
    suffix = target.suffix
    for index in range(2, 15):
        candidate = target.with_name(f"{stem}_v{index}{suffix}")
        if not candidate.exists():
            return candidate
    return target.with_name(f"{stem}_latest{suffix}")


def make_canvas(title: str, subtitle: str = "", size=(1800, 1100)):
    img = Image.new("RGB", size, PAGE_BG)
    draw = ImageDraw.Draw(img)
    draw_centered(draw, (size[0] // 2, 48), title, FONT_TITLE, TITLE)
    if subtitle:
        draw_centered(draw, (size[0] // 2, 90), subtitle, FONT_SUBTITLE, MUTED)
    return img, draw


def text_size(draw, text, font):
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw, text, font, max_width):
    raw_lines = str(text).split("\n")
    wrapped = []
    for raw in raw_lines:
        words = raw.split()
        if not words:
            wrapped.append("")
            continue
        current = words[0]
        for word in words[1:]:
            trial = f"{current} {word}"
            width, _ = text_size(draw, trial, font)
            if width <= max_width:
                current = trial
            else:
                wrapped.append(current)
                current = word
        wrapped.append(current)
    return wrapped


def draw_centered(draw, xy, text, font, fill):
    width, height = text_size(draw, text, font)
    draw.text((xy[0] - width / 2, xy[1] - height / 2), text, font=font, fill=fill)


def draw_multiline(draw, box, text, font, fill, align="center", line_spacing=5, padding=10):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(20, x2 - x1 - padding * 2))
    sizes = [text_size(draw, line, font) for line in lines]
    total_height = sum(height for _, height in sizes) + line_spacing * (len(lines) - 1 if lines else 0)
    y = y1 + ((y2 - y1 - total_height) / 2)
    for line, (width, height) in zip(lines, sizes):
        if align == "left":
            x = x1 + padding
        elif align == "right":
            x = x2 - width - padding
        else:
            x = x1 + ((x2 - x1 - width) / 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += height + line_spacing


def draw_round_box(draw, rect, title, body="", fill=SLATE_FILL, outline=LINE):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline=outline, width=3)
    if body:
        title_box = (x1 + 8, y1 + 12, x2 - 8, y1 + 48)
        body_box = (x1 + 10, y1 + 55, x2 - 10, y2 - 10)
        draw_multiline(draw, title_box, title, FONT_BOLD, TEXT)
        draw_multiline(draw, body_box, body, FONT_SMALL, TEXT)
    else:
        draw_multiline(draw, rect, title, FONT_BOLD, TEXT)


def draw_actor(draw, x, y, label):
    draw.ellipse((x - 18, y, x + 18, y + 36), outline=TEXT, width=3)
    draw.line((x, y + 36, x, y + 100), fill=TEXT, width=3)
    draw.line((x - 34, y + 56, x + 34, y + 56), fill=TEXT, width=3)
    draw.line((x, y + 100, x - 28, y + 145), fill=TEXT, width=3)
    draw.line((x, y + 100, x + 28, y + 145), fill=TEXT, width=3)
    draw_multiline(draw, (x - 70, y + 152, x + 70, y + 210), label, FONT_SMALL_BOLD, TEXT)


def draw_usecase(draw, rect, label, fill=BLUE_FILL, outline=BLUE):
    draw.ellipse(rect, fill=fill, outline=outline, width=3)
    draw_multiline(draw, rect, label, FONT_SMALL_BOLD, TEXT)


def center_of(rect):
    x1, y1, x2, y2 = rect
    return ((x1 + x2) / 2, (y1 + y2) / 2)


def draw_dashed_line(draw, start, end, fill=LINE, dash=10, gap=6, width=2):
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    distance = max(1, int((dx * dx + dy * dy) ** 0.5))
    for step in range(0, distance, dash + gap):
        start_ratio = step / distance
        end_ratio = min(distance, step + dash) / distance
        sx = x1 + dx * start_ratio
        sy = y1 + dy * start_ratio
        ex = x1 + dx * end_ratio
        ey = y1 + dy * end_ratio
        draw.line((sx, sy, ex, ey), fill=fill, width=width)


def draw_arrow(draw, start, end, fill=BLUE, dashed=False, width=3):
    if dashed:
        draw_dashed_line(draw, start, end, fill=fill, width=max(1, width - 1))
    else:
        draw.line((start, end), fill=fill, width=width)

    dx = end[0] - start[0]
    dy = end[1] - start[1]
    distance = max(1.0, (dx * dx + dy * dy) ** 0.5)
    ux = dx / distance
    uy = dy / distance
    left = (
        end[0] - 18 * ux - 8 * uy,
        end[1] - 18 * uy + 8 * ux,
    )
    right = (
        end[0] - 18 * ux + 8 * uy,
        end[1] - 18 * uy - 8 * ux,
    )
    draw.polygon([end, left, right], fill=fill)


def draw_connector(draw, start, end, label="", color=LINE, dashed=False, arrow=False, label_offset=-24):
    if arrow:
        draw_arrow(draw, start, end, fill=color, dashed=dashed, width=3)
    else:
        if dashed:
            draw_dashed_line(draw, start, end, fill=color, width=2)
        else:
            draw.line((start, end), fill=color, width=3)
    if label:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2 + label_offset
        width, height = text_size(draw, label, FONT_TINY)
        draw.rounded_rectangle(
            (mx - width / 2 - 8, my - 4, mx + width / 2 + 8, my + height + 4),
            radius=8,
            fill=PAGE_BG,
            outline=None,
        )
        draw.text((mx - width / 2, my), label, font=FONT_TINY, fill=MUTED)


def anchor_point(rect, toward_rect):
    sx, sy = center_of(rect)
    tx, ty = center_of(toward_rect)
    x1, y1, x2, y2 = rect
    dx = tx - sx
    dy = ty - sy
    if abs(dx) >= abs(dy):
        return (x2, sy) if dx >= 0 else (x1, sy)
    return (sx, y2) if dy >= 0 else (sx, y1)


def draw_note(draw, x, y, text, width=360):
    lines = wrap_text(draw, text, FONT_TINY, width - 24)
    line_heights = [text_size(draw, line, FONT_TINY)[1] for line in lines]
    height = sum(line_heights) + (len(lines) - 1) * 4 + 20
    rect = (x, y, x + width, y + height)
    draw.rounded_rectangle(rect, radius=14, fill=NOTE_FILL, outline=AMBER, width=2)
    draw_multiline(draw, rect, text, FONT_TINY, TEXT, align="left", padding=12)
    return rect


def draw_class_box(draw, rect, name, attrs, ops, fill=SLATE_FILL, outline=SLATE):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, radius=16, fill=fill, outline=outline, width=3)
    header_bottom = y1 + 52
    draw.rectangle((x1, y1, x2, header_bottom), fill=outline)
    draw_multiline(draw, (x1 + 8, y1 + 8, x2 - 8, header_bottom - 8), name, FONT_SMALL_BOLD, PAGE_BG)

    attrs_height = max(80, min(180, 28 * max(2, len(attrs))))
    split_y = header_bottom + attrs_height
    draw.line((x1, header_bottom, x2, header_bottom), fill=outline, width=2)
    draw.line((x1, split_y, x2, split_y), fill=outline, width=2)

    attr_y = header_bottom + 12
    for attr in attrs:
        draw.text((x1 + 14, attr_y), attr, font=FONT_TINY, fill=TEXT)
        attr_y += 24

    op_y = split_y + 12
    for op in ops:
        draw.text((x1 + 14, op_y), op, font=FONT_TINY, fill=TEXT)
        op_y += 24


def draw_participant(draw, x, top, label):
    rect = (x - 90, top, x + 90, top + 48)
    draw.rounded_rectangle(rect, radius=12, fill=BLUE_FILL, outline=BLUE, width=2)
    draw_multiline(draw, rect, label, FONT_SMALL_BOLD, TEXT)
    draw_dashed_line(draw, (x, top + 48), (x, top + 820), fill=SLATE, dash=12, gap=10, width=2)
    return rect


def draw_sequence_message(draw, x1, x2, y, label, returned=False):
    color = SLATE if returned else BLUE
    draw_arrow(draw, (x1, y), (x2, y), fill=color, dashed=returned, width=3)
    label_box = (
        min(x1, x2) + 10,
        y - 34,
        max(x1, x2) - 10,
        y - 8,
    )
    draw_multiline(draw, label_box, label, FONT_TINY, TEXT, line_spacing=2)


def use_case_specs():
    return [
        {
            "slug": "usecase_01_vue_globale",
            "title": "Diagramme de cas d'utilisation - Vue globale",
            "description": "Vue d'ensemble des acteurs et des fonctions majeures de SmartArchive.",
            "actors": [
                {"id": "visitor", "label": "Visiteur", "x": 110, "y": 250},
                {"id": "user", "label": "Utilisateur", "x": 110, "y": 610},
                {"id": "admin", "label": "Administrateur", "x": 1690, "y": 430},
            ],
            "usecases": [
                {"id": "cin_upload", "label": "Uploader image\nCIN", "rect": (370, 170, 690, 265)},
                {"id": "email_code", "label": "Saisir email\net code", "rect": (370, 310, 690, 405)},
                {"id": "login", "label": "Se connecter", "rect": (370, 450, 690, 545)},
                {"id": "reset_pw", "label": "Réinitialiser\nmot de passe", "rect": (370, 590, 690, 685)},
                {"id": "profile", "label": "Gérer profil", "rect": (780, 170, 1100, 265)},
                {"id": "upload_doc", "label": "Uploader\nun document", "rect": (780, 310, 1100, 405)},
                {"id": "history", "label": "Consulter\nhistorique", "rect": (780, 450, 1100, 545)},
                {"id": "my_cin", "label": "Consulter\nma CIN", "rect": (780, 590, 1100, 685)},
                {"id": "correct", "label": "Corriger\nl'extraction", "rect": (780, 730, 1100, 825)},
                {"id": "users", "label": "Gérer\nutilisateurs", "rect": (1190, 170, 1510, 265)},
                {"id": "review", "label": "Réviser\nles documents", "rect": (1190, 310, 1510, 405)},
                {"id": "stats", "label": "Consulter\nstatistiques", "rect": (1190, 450, 1510, 545)},
                {"id": "reports", "label": "Générer\nrapports", "rect": (1190, 590, 1510, 685)},
                {"id": "settings", "label": "Gérer\nparamètres", "rect": (1190, 730, 1510, 825)},
            ],
            "links": [
                ("visitor", "cin_upload", "assoc"),
                ("visitor", "email_code", "assoc"),
                ("visitor", "login", "assoc"),
                ("visitor", "reset_pw", "assoc"),
                ("user", "profile", "assoc"),
                ("user", "upload_doc", "assoc"),
                ("user", "history", "assoc"),
                ("user", "my_cin", "assoc"),
                ("user", "correct", "assoc"),
                ("admin", "users", "assoc"),
                ("admin", "review", "assoc"),
                ("admin", "stats", "assoc"),
                ("admin", "reports", "assoc"),
                ("admin", "settings", "assoc"),
            ],
        },
        {
            "slug": "usecase_02_authentification_onboarding",
            "title": "Diagramme de cas d'utilisation - Authentification et onboarding",
            "description": "Scénarios d'entrée: CIN, code email, vérification et récupération du mot de passe.",
            "actors": [
                {"id": "visitor", "label": "Visiteur", "x": 110, "y": 340},
                {"id": "mail", "label": "Service\nEmail", "x": 1685, "y": 280},
                {"id": "ocr", "label": "OCR CIN\nPython", "x": 1685, "y": 600},
            ],
            "usecases": [
                {"id": "scan_cin", "label": "Déposer image\nCIN", "rect": (380, 200, 720, 300)},
                {"id": "extract_cin", "label": "Extraire CIN\net identité", "rect": (790, 200, 1130, 300)},
                {"id": "check_account", "label": "Détecter compte\nexistant / nouveau", "rect": (1200, 200, 1540, 300)},
                {"id": "enter_email", "label": "Saisir / confirmer\nemail", "rect": (380, 400, 720, 500)},
                {"id": "send_code", "label": "Envoyer code\nde vérification", "rect": (790, 400, 1130, 500)},
                {"id": "verify_code", "label": "Vérifier le code", "rect": (1200, 400, 1540, 500)},
                {"id": "activate_account", "label": "Activer le compte\net ouvrir la session", "rect": (790, 620, 1130, 720)},
                {"id": "forgot_password", "label": "Lancer la\nrécupération", "rect": (380, 620, 720, 720)},
                {"id": "reset_password", "label": "Réinitialiser\nmot de passe", "rect": (1200, 620, 1540, 720)},
            ],
            "links": [
                ("visitor", "scan_cin", "assoc"),
                ("visitor", "enter_email", "assoc"),
                ("visitor", "verify_code", "assoc"),
                ("visitor", "forgot_password", "assoc"),
                ("ocr", "extract_cin", "assoc"),
                ("mail", "send_code", "assoc"),
                ("scan_cin", "extract_cin", "include"),
                ("extract_cin", "check_account", "include"),
                ("enter_email", "send_code", "include"),
                ("send_code", "verify_code", "include"),
                ("verify_code", "activate_account", "include"),
                ("forgot_password", "send_code", "include"),
                ("forgot_password", "reset_password", "extend"),
            ],
        },
        {
            "slug": "usecase_03_espace_utilisateur",
            "title": "Diagramme de cas d'utilisation - Espace utilisateur",
            "description": "Fonctions offertes au compte utilisateur authentifié.",
            "actors": [
                {"id": "user", "label": "Utilisateur", "x": 120, "y": 430},
                {"id": "ocr", "label": "OCR /\nExtractor", "x": 1675, "y": 250},
            ],
            "usecases": [
                {"id": "dash", "label": "Consulter\nle dashboard", "rect": (400, 160, 760, 255)},
                {"id": "user_profile", "label": "Modifier profil /\nchanger mot de passe", "rect": (400, 300, 760, 395)},
                {"id": "user_upload", "label": "Téléverser\nun document", "rect": (400, 440, 760, 535)},
                {"id": "user_history", "label": "Consulter\nhistorique", "rect": (400, 580, 760, 675)},
                {"id": "user_details", "label": "Voir détail\ndocument", "rect": (840, 160, 1200, 255)},
                {"id": "user_download", "label": "Télécharger\nle résultat", "rect": (840, 300, 1200, 395)},
                {"id": "user_edit", "label": "Corriger\nentités et texte", "rect": (840, 440, 1200, 535)},
                {"id": "user_delete", "label": "Supprimer\ndocument", "rect": (840, 580, 1200, 675)},
                {"id": "user_cin", "label": "Afficher\nma CIN", "rect": (1280, 250, 1590, 345)},
                {"id": "user_notifications", "label": "Consulter\nnotifications", "rect": (1280, 460, 1590, 555)},
            ],
            "links": [
                ("user", "dash", "assoc"),
                ("user", "user_profile", "assoc"),
                ("user", "user_upload", "assoc"),
                ("user", "user_history", "assoc"),
                ("user", "user_details", "assoc"),
                ("user", "user_download", "assoc"),
                ("user", "user_edit", "assoc"),
                ("user", "user_delete", "assoc"),
                ("user", "user_cin", "assoc"),
                ("user", "user_notifications", "assoc"),
                ("ocr", "user_upload", "assoc"),
                ("user_history", "user_details", "include"),
                ("user_details", "user_download", "extend"),
                ("user_details", "user_edit", "extend"),
                ("user_history", "user_delete", "extend"),
            ],
        },
        {
            "slug": "usecase_04_espace_admin",
            "title": "Diagramme de cas d'utilisation - Espace administrateur",
            "description": "Fonctions de pilotage, contrôle et maintenance réservées à l'administrateur.",
            "actors": [
                {"id": "admin", "label": "Administrateur", "x": 110, "y": 430},
                {"id": "ia", "label": "Moteur IA /\nRanking", "x": 1685, "y": 400},
            ],
            "usecases": [
                {"id": "admin_dash", "label": "Consulter\nle dashboard admin", "rect": (380, 150, 760, 245)},
                {"id": "list_users", "label": "Lister les\nutilisateurs", "rect": (380, 290, 760, 385)},
                {"id": "edit_user", "label": "Modifier / vérifier\nun utilisateur", "rect": (380, 430, 760, 525)},
                {"id": "delete_user", "label": "Supprimer\nun utilisateur", "rect": (380, 570, 760, 665)},
                {"id": "list_docs", "label": "Lister les\ndocuments", "rect": (840, 150, 1220, 245)},
                {"id": "review_doc", "label": "Réviser et valider\nun document", "rect": (840, 290, 1220, 385)},
                {"id": "admin_stats", "label": "Consulter\nstatistiques", "rect": (840, 430, 1220, 525)},
                {"id": "admin_reports", "label": "Produire les\nrapports", "rect": (840, 570, 1220, 665)},
                {"id": "admin_settings", "label": "Mettre à jour\nles paramètres", "rect": (1280, 290, 1590, 385)},
                {"id": "admin_auto", "label": "Obtenir des suggestions\net autocomplétion", "rect": (1280, 500, 1590, 595)},
            ],
            "links": [
                ("admin", "admin_dash", "assoc"),
                ("admin", "list_users", "assoc"),
                ("admin", "edit_user", "assoc"),
                ("admin", "delete_user", "assoc"),
                ("admin", "list_docs", "assoc"),
                ("admin", "review_doc", "assoc"),
                ("admin", "admin_stats", "assoc"),
                ("admin", "admin_reports", "assoc"),
                ("admin", "admin_settings", "assoc"),
                ("admin", "admin_auto", "assoc"),
                ("ia", "admin_auto", "assoc"),
                ("list_users", "edit_user", "include"),
                ("list_users", "delete_user", "extend"),
                ("list_docs", "review_doc", "include"),
                ("admin_reports", "admin_auto", "include"),
            ],
        },
        {
            "slug": "usecase_05_reporting",
            "title": "Diagramme de cas d'utilisation - Rapports et analytique",
            "description": "Scénarios autour des suggestions, filtres et exports dans le module reporting.",
            "actors": [
                {"id": "admin", "label": "Administrateur", "x": 120, "y": 430},
                {"id": "ia", "label": "Service Python\nOllama / JS ranker", "x": 1670, "y": 280},
                {"id": "db", "label": "Base Users /\nDocuments", "x": 1670, "y": 610},
            ],
            "usecases": [
                {"id": "filter_name", "label": "Filtrer par\nnom", "rect": (420, 170, 760, 265)},
                {"id": "filter_cin", "label": "Filtrer par\nCIN", "rect": (420, 330, 760, 425)},
                {"id": "get_suggestions", "label": "Obtenir suggestions\nnom / CIN", "rect": (840, 170, 1180, 265)},
                {"id": "text_auto", "label": "Autocomplétion\nde texte", "rect": (840, 330, 1180, 425)},
                {"id": "token_auto", "label": "Autocomplétion\nde tokens", "rect": (840, 490, 1180, 585)},
                {"id": "load_rows", "label": "Charger les lignes\nde rapport", "rect": (1260, 170, 1540, 265)},
                {"id": "export_draft", "label": "Exporter le\nbrouillon", "rect": (1260, 390, 1540, 485)},
            ],
            "links": [
                ("admin", "filter_name", "assoc"),
                ("admin", "filter_cin", "assoc"),
                ("admin", "get_suggestions", "assoc"),
                ("admin", "text_auto", "assoc"),
                ("admin", "token_auto", "assoc"),
                ("admin", "load_rows", "assoc"),
                ("admin", "export_draft", "assoc"),
                ("ia", "get_suggestions", "assoc"),
                ("ia", "text_auto", "assoc"),
                ("ia", "token_auto", "assoc"),
                ("db", "get_suggestions", "assoc"),
                ("db", "load_rows", "assoc"),
                ("filter_name", "load_rows", "include"),
                ("filter_cin", "load_rows", "include"),
            ],
        },
    ]


def class_specs():
    return [
        {
            "slug": "class_01_entites_globales",
            "title": "Diagramme de classes - Entités globales",
            "description": "Les principales entités persistantes manipulées par SmartArchive.",
            "classes": [
                {
                    "id": "account",
                    "name": "<<abstract>> Account",
                    "rect": (710, 120, 1110, 390),
                    "fill": BLUE_FILL,
                    "outline": BLUE,
                    "attrs": [
                        "- email: String",
                        "- password: String",
                        "- username: String",
                        "- cin_number: String",
                        "- is_verified: Boolean",
                    ],
                    "ops": [
                        "+ login()",
                        "+ updateProfile()",
                        "+ changePassword()",
                    ],
                },
                {
                    "id": "user",
                    "name": "User",
                    "rect": (210, 500, 610, 830),
                    "fill": GREEN_FILL,
                    "outline": GREEN,
                    "attrs": [
                        "- first_name: String",
                        "- last_name: String",
                        "- full_name: String",
                        "- phone: String",
                        "- address: String",
                        "- last_login: Date",
                    ],
                    "ops": [
                        "+ verifyEmail()",
                        "+ getHistory()",
                        "+ getMyCIN()",
                    ],
                },
                {
                    "id": "admin",
                    "name": "Admin",
                    "rect": (1210, 500, 1610, 830),
                    "fill": ROSE_FILL,
                    "outline": ROSE,
                    "attrs": [
                        "- role: 'admin'",
                        "- last_login: Date",
                        "- phone: String",
                        "- address: String",
                    ],
                    "ops": [
                        "+ reviewDocument()",
                        "+ updateUser()",
                        "+ updateSettings()",
                    ],
                },
                {
                    "id": "verification",
                    "name": "VerificationCode",
                    "rect": (700, 500, 1120, 780),
                    "fill": AMBER_FILL,
                    "outline": AMBER,
                    "attrs": [
                        "- email: String",
                        "- code: String",
                        "- purpose: verification|reset",
                        "- expires_at: Date",
                    ],
                    "ops": [
                        "+ generateCode()",
                        "+ validate()",
                        "+ deleteOnSuccess()",
                    ],
                },
                {
                    "id": "document",
                    "name": "DocumentRecord",
                    "rect": (140, 870, 960, 1160),
                    "fill": VIOLET_FILL,
                    "outline": VIOLET,
                    "attrs": [
                        "- filename: String",
                        "- type: String",
                        "- status: String",
                        "- verification_status: String",
                        "- quality_score: Number",
                        "- extracted_data: Mixed",
                        "- full_text: String",
                        "- requires_admin_review: Boolean",
                    ],
                    "ops": [
                        "+ save()",
                        "+ updateEntities()",
                        "+ delete()",
                    ],
                },
                {
                    "id": "settings",
                    "name": "AdminSettings",
                    "rect": (1080, 900, 1600, 1160),
                    "fill": SLATE_FILL,
                    "outline": SLATE,
                    "attrs": [
                        "- siteName: String",
                        "- locale: String",
                        "- timezone: String",
                        "- enforce2FA: Boolean",
                        "- autoVerifyThreshold: Number",
                    ],
                    "ops": [
                        "+ loadDefault()",
                        "+ update()",
                    ],
                },
            ],
            "relations": [
                ("user", "account", "inheritance", "hérite"),
                ("admin", "account", "inheritance", "hérite"),
                ("verification", "user", "dependency", "vérifie"),
                ("verification", "admin", "dependency", "vérifie"),
                ("user", "document", "association", "possède 0..*"),
                ("admin", "document", "association", "révise 0..*"),
                ("admin", "settings", "association", "configure"),
            ],
        },
        {
            "slug": "class_02_contexte_auth",
            "title": "Diagramme de classes - Contexte authentification",
            "description": "Les composants logiques impliqués dans la vérification d'identité et l'ouverture de session.",
            "classes": [
                {
                    "id": "auth_context",
                    "name": "AuthContext (frontend)",
                    "rect": (110, 180, 470, 470),
                    "fill": BLUE_FILL,
                    "outline": BLUE,
                    "attrs": [
                        "- user: Object",
                        "- isAdmin: Boolean",
                        "- loading: Boolean",
                    ],
                    "ops": [
                        "+ loadUser()",
                        "+ login()",
                        "+ verifyEmail()",
                        "+ resendCode()",
                    ],
                },
                {
                    "id": "auth_controller",
                    "name": "AuthRoutes / Controller",
                    "rect": (610, 120, 1030, 470),
                    "fill": GREEN_FILL,
                    "outline": GREEN,
                    "attrs": [
                        "- /login",
                        "- /verify-email",
                        "- /send-code",
                        "- /verify-code",
                        "- /change-password",
                    ],
                    "ops": [
                        "+ handleLogin()",
                        "+ handleVerifyEmail()",
                        "+ handleResetPassword()",
                    ],
                },
                {
                    "id": "email_service",
                    "name": "EmailService",
                    "rect": (1170, 120, 1540, 400),
                    "fill": AMBER_FILL,
                    "outline": AMBER,
                    "attrs": [
                        "- SMTP config",
                        "- welcome templates",
                        "- verification templates",
                    ],
                    "ops": [
                        "+ sendVerificationEmail()",
                        "+ sendWelcomeEmail()",
                        "+ sendWelcomeEmailToNewUser()",
                    ],
                },
                {
                    "id": "token_service",
                    "name": "TokenHelper",
                    "rect": (1170, 470, 1540, 700),
                    "fill": VIOLET_FILL,
                    "outline": VIOLET,
                    "attrs": [
                        "- PASSWORD_SECRET",
                        "- JWT payload",
                    ],
                    "ops": [
                        "+ generateToken()",
                        "+ generateVerificationCode()",
                        "+ isStrongPassword()",
                    ],
                },
                {
                    "id": "verify_code",
                    "name": "VerificationCode",
                    "rect": (610, 560, 1030, 860),
                    "fill": ROSE_FILL,
                    "outline": ROSE,
                    "attrs": [
                        "- email",
                        "- code",
                        "- purpose",
                        "- expires_at",
                    ],
                    "ops": [
                        "+ create()",
                        "+ findOne()",
                        "+ deleteOne()",
                    ],
                },
                {
                    "id": "identity_models",
                    "name": "User / Admin",
                    "rect": (120, 610, 470, 880),
                    "fill": SLATE_FILL,
                    "outline": SLATE,
                    "attrs": [
                        "- cin_number",
                        "- email",
                        "- username",
                        "- is_verified",
                    ],
                    "ops": [
                        "+ findOne()",
                        "+ save()",
                        "+ updateOne()",
                    ],
                },
            ],
            "relations": [
                ("auth_context", "auth_controller", "association", "consomme"),
                ("auth_controller", "identity_models", "association", "charge"),
                ("auth_controller", "verify_code", "association", "lit/écrit"),
                ("auth_controller", "email_service", "association", "envoie"),
                ("auth_controller", "token_service", "association", "génère"),
            ],
        },
        {
            "slug": "class_03_contexte_document",
            "title": "Diagramme de classes - Contexte documentaire",
            "description": "Les composants du pipeline OCR, extraction, classification et correction documentaire.",
            "classes": [
                {
                    "id": "upload_ui",
                    "name": "UploadDocument (frontend)",
                    "rect": (120, 170, 500, 450),
                    "fill": BLUE_FILL,
                    "outline": BLUE,
                    "attrs": [
                        "- selectedFile: File",
                        "- documentType: String",
                        "- progress: Number",
                    ],
                    "ops": [
                        "+ handleUpload()",
                        "+ submitFormData()",
                    ],
                },
                {
                    "id": "document_api",
                    "name": "UploadRoutes / Controller",
                    "rect": (620, 120, 1080, 500),
                    "fill": GREEN_FILL,
                    "outline": GREEN,
                    "attrs": [
                        "- /documents/upload",
                        "- /documents/history",
                        "- /documents/:id",
                        "- /documents/:id/entities",
                        "- /documents/:id",
                    ],
                    "ops": [
                        "+ classifyDocument()",
                        "+ computeQualityScore()",
                        "+ canAccessDocument()",
                        "+ saveToMongoDB()",
                    ],
                },
                {
                    "id": "ocr_service",
                    "name": "OCRService",
                    "rect": (1210, 120, 1610, 360),
                    "fill": AMBER_FILL,
                    "outline": AMBER,
                    "attrs": [
                        "- OCR_API_URL",
                        "- /ocr",
                    ],
                    "ops": [
                        "+ extractText()",
                        "+ fallbackMock()",
                    ],
                },
                {
                    "id": "extractor",
                    "name": "EntityExtractor",
                    "rect": (1210, 430, 1610, 670),
                    "fill": VIOLET_FILL,
                    "outline": VIOLET,
                    "attrs": [
                        "- EXTRACTOR_API_URL",
                        "- /extract",
                    ],
                    "ops": [
                        "+ extractEntities()",
                        "+ fallbackMock()",
                    ],
                },
                {
                    "id": "record",
                    "name": "DocumentRecord",
                    "rect": (600, 620, 1100, 980),
                    "fill": ROSE_FILL,
                    "outline": ROSE,
                    "attrs": [
                        "- filename",
                        "- status",
                        "- quality_score",
                        "- requires_admin_review",
                        "- extracted_data",
                        "- full_text",
                    ],
                    "ops": [
                        "+ save()",
                        "+ find()",
                        "+ findById()",
                        "+ findByIdAndDelete()",
                    ],
                },
                {
                    "id": "history_ui",
                    "name": "DocumentHistory / Details",
                    "rect": (120, 620, 500, 980),
                    "fill": SLATE_FILL,
                    "outline": SLATE,
                    "attrs": [
                        "- filters",
                        "- viewingDoc",
                        "- edited_data",
                    ],
                    "ops": [
                        "+ getHistory()",
                        "+ getById()",
                        "+ updateEntities()",
                        "+ delete()",
                    ],
                },
            ],
            "relations": [
                ("upload_ui", "document_api", "association", "envoie"),
                ("history_ui", "document_api", "association", "consomme"),
                ("document_api", "ocr_service", "association", "appelle"),
                ("document_api", "extractor", "association", "appelle"),
                ("document_api", "record", "association", "persiste"),
            ],
        },
        {
            "slug": "class_04_contexte_admin_reporting",
            "title": "Diagramme de classes - Contexte administration et reporting",
            "description": "Les classes logiques mobilisées pour la supervision, le ranking et les paramètres.",
            "classes": [
                {
                    "id": "admin_ui",
                    "name": "Admin UI Modules",
                    "rect": (120, 150, 470, 470),
                    "fill": BLUE_FILL,
                    "outline": BLUE,
                    "attrs": [
                        "- UsersList",
                        "- DocumentDetails",
                        "- Reports",
                        "- Settings",
                        "- Statistics",
                    ],
                    "ops": [
                        "+ getAllUsers()",
                        "+ reviewDocument()",
                        "+ getReports()",
                        "+ updateSettings()",
                    ],
                },
                {
                    "id": "admin_controller",
                    "name": "AdminRoutes / Controller",
                    "rect": (580, 110, 1040, 510),
                    "fill": GREEN_FILL,
                    "outline": GREEN,
                    "attrs": [
                        "- /users",
                        "- /statistics",
                        "- /documents/:id/review",
                        "- /reports/*",
                        "- /settings",
                    ],
                    "ops": [
                        "+ ensureAdmin()",
                        "+ rankSuggestions()",
                        "+ buildUniqueNameCinPairs()",
                    ],
                },
                {
                    "id": "ranker",
                    "name": "ReportRankingService",
                    "rect": (1160, 120, 1580, 390),
                    "fill": AMBER_FILL,
                    "outline": AMBER,
                    "attrs": [
                        "- ollama_autocomplete.py",
                        "- jsRankSuggestions()",
                    ],
                    "ops": [
                        "+ runPythonAutocomplete()",
                        "+ rankSuggestions()",
                    ],
                },
                {
                    "id": "stats_service",
                    "name": "StatisticsService",
                    "rect": (1160, 450, 1580, 680),
                    "fill": VIOLET_FILL,
                    "outline": VIOLET,
                    "attrs": [
                        "- totalUsers",
                        "- verifiedUsers",
                        "- totalDocuments",
                        "- documentsToday",
                    ],
                    "ops": [
                        "+ aggregateKPIs()",
                        "+ buildMonthlySeries()",
                    ],
                },
                {
                    "id": "data_models",
                    "name": "User + DocumentRecord",
                    "rect": (560, 660, 1040, 980),
                    "fill": ROSE_FILL,
                    "outline": ROSE,
                    "attrs": [
                        "- username",
                        "- cin_number",
                        "- status",
                        "- quality_score",
                    ],
                    "ops": [
                        "+ find()",
                        "+ countDocuments()",
                        "+ findByIdAndUpdate()",
                    ],
                },
                {
                    "id": "settings",
                    "name": "AdminSettings",
                    "rect": (120, 630, 470, 930),
                    "fill": SLATE_FILL,
                    "outline": SLATE,
                    "attrs": [
                        "- locale",
                        "- timezone",
                        "- sessionTimeout",
                        "- autoVerifyThreshold",
                    ],
                    "ops": [
                        "+ createDefault()",
                        "+ update()",
                    ],
                },
            ],
            "relations": [
                ("admin_ui", "admin_controller", "association", "consomme"),
                ("admin_controller", "ranker", "association", "appelle"),
                ("admin_controller", "stats_service", "association", "utilise"),
                ("admin_controller", "data_models", "association", "lit/écrit"),
                ("admin_controller", "settings", "association", "met à jour"),
            ],
        },
    ]


def sequence_specs():
    return [
        {
            "slug": "sequence_01_upload_cin_detection_compte",
            "title": "Diagramme de séquence - Upload CIN et détection du compte",
            "description": "Scénario principal depuis l'écran Auth jusqu'au retour user_exists / admin_exists / new_user.",
            "participants": ["Visiteur", "Interface Auth", "API CIN", "OCR Python", "MongoDB"],
            "messages": [
                {"kind": "section", "label": "Flux principal"},
                {"kind": "call", "from": 0, "to": 1, "label": "Déposer image CIN"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /cin/upload(file)"},
                {"kind": "call", "from": 2, "to": 3, "label": "Traiter image /process"},
                {"kind": "return", "from": 3, "to": 2, "label": "Texte + champs extraits"},
                {"kind": "call", "from": 2, "to": 4, "label": "Chercher User/Admin + upsert DocumentRecord"},
                {"kind": "return", "from": 4, "to": 2, "label": "Statut du compte"},
                {"kind": "return", "from": 2, "to": 1, "label": "cin_validation + account flags"},
                {"kind": "return", "from": 1, "to": 0, "label": "Afficher branche adaptée"},
                {"kind": "note", "participant": 2, "label": "Branches: CIN invalide | admin existant | utilisateur vérifié | utilisateur non vérifié | nouveau compte"},
            ],
        },
        {
            "slug": "sequence_02_envoi_code_verification",
            "title": "Diagramme de séquence - Envoi du code de vérification",
            "description": "Scénario /api/auth/send-code pour nouvel utilisateur, utilisateur existant ou admin.",
            "participants": ["Interface Auth", "API Auth", "Users/Admins DB", "VerificationCode DB", "Service Email"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "POST /auth/send-code(email, purpose, cin_number)"},
                {"kind": "call", "from": 1, "to": 2, "label": "Vérifier email / CIN existants"},
                {"kind": "return", "from": 2, "to": 1, "label": "Compte trouvé ou à créer"},
                {"kind": "call", "from": 1, "to": 3, "label": "Créer code + expiration"},
                {"kind": "call", "from": 1, "to": 4, "label": "Envoyer email de vérification"},
                {"kind": "return", "from": 4, "to": 1, "label": "Envoi SMTP OK"},
                {"kind": "return", "from": 1, "to": 0, "label": "Success + email + purpose"},
                {"kind": "note", "participant": 1, "label": "Cas alternatif: email déjà lié à un autre compte lors d'un nouveau flux d'inscription"},
            ],
        },
        {
            "slug": "sequence_03_verification_email_code",
            "title": "Diagramme de séquence - Vérification email et activation",
            "description": "Scénario /api/auth/verify-email après saisie du code à 6 chiffres.",
            "participants": ["Utilisateur", "Écran VerifyEmail", "API Auth", "VerificationCode DB", "Users/Admins DB", "Service Email"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Saisir le code"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /auth/verify-email(email, code)"},
                {"kind": "call", "from": 2, "to": 3, "label": "Chercher code valide non expiré"},
                {"kind": "return", "from": 3, "to": 2, "label": "Code trouvé"},
                {"kind": "call", "from": 2, "to": 4, "label": "Charger User ou Admin"},
                {"kind": "return", "from": 4, "to": 2, "label": "Compte correspondant"},
                {"kind": "call", "from": 2, "to": 4, "label": "Marquer is_verified si nécessaire"},
                {"kind": "call", "from": 2, "to": 5, "label": "Envoyer welcome email"},
                {"kind": "call", "from": 2, "to": 3, "label": "Supprimer code consommé"},
                {"kind": "return", "from": 2, "to": 1, "label": "token + user"},
                {"kind": "return", "from": 1, "to": 0, "label": "Navigation vers /welcom"},
                {"kind": "note", "participant": 2, "label": "Cas alternatif: code invalide ou expiré -> erreur de vérification"},
            ],
        },
        {
            "slug": "sequence_04_login",
            "title": "Diagramme de séquence - Connexion classique",
            "description": "Scénario /api/auth/login avec recherche prioritaire d'un admin puis d'un utilisateur.",
            "participants": ["Utilisateur/Admin", "Form Login", "API Auth", "Users/Admins DB", "Service Email"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Saisir CIN + mot de passe"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /auth/login(cin_number, password)"},
                {"kind": "call", "from": 2, "to": 3, "label": "Chercher Admin puis User"},
                {"kind": "return", "from": 3, "to": 2, "label": "Compte trouvé ou non"},
                {"kind": "call", "from": 2, "to": 4, "label": "Envoyer welcome email si succès"},
                {"kind": "return", "from": 2, "to": 1, "label": "token + user"},
                {"kind": "return", "from": 1, "to": 0, "label": "Redirection حسب le rôle"},
                {"kind": "note", "participant": 2, "label": "Cas alternatifs: carte introuvable | mot de passe incorrect | user non vérifié activé au passage"},
            ],
        },
        {
            "slug": "sequence_05_reset_password",
            "title": "Diagramme de séquence - Récupération et réinitialisation du mot de passe",
            "description": "Scénario ForgotPassword: envoi du code puis mise à jour du nouveau mot de passe.",
            "participants": ["Utilisateur", "ForgotPassword UI", "API Auth", "VerificationCode DB", "User DB"],
            "messages": [
                {"kind": "section", "label": "Étape 1 - Envoi du code"},
                {"kind": "call", "from": 0, "to": 1, "label": "Saisir email"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /auth/send-code(purpose=reset)"},
                {"kind": "call", "from": 2, "to": 3, "label": "Créer code password_reset"},
                {"kind": "return", "from": 2, "to": 1, "label": "Code envoyé"},
                {"kind": "section", "label": "Étape 2 - Réinitialisation"},
                {"kind": "call", "from": 0, "to": 1, "label": "Saisir code + nouveau mot de passe"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /auth/verify-code"},
                {"kind": "call", "from": 2, "to": 3, "label": "Vérifier code non expiré"},
                {"kind": "return", "from": 3, "to": 2, "label": "Code valide"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /auth/reset-password"},
                {"kind": "call", "from": 2, "to": 4, "label": "Mettre à jour password chiffré"},
                {"kind": "call", "from": 2, "to": 3, "label": "Supprimer code"},
                {"kind": "return", "from": 2, "to": 1, "label": "Password reset success"},
            ],
        },
        {
            "slug": "sequence_06_dashboard_utilisateur",
            "title": "Diagramme de séquence - Chargement du dashboard utilisateur",
            "description": "Le tableau de bord utilisateur combine l'historique et les statistiques personnelles.",
            "participants": ["Utilisateur", "UserDashboard UI", "API Documents", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /dashboard"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/history"},
                {"kind": "call", "from": 2, "to": 3, "label": "Lister documents du CIN / username"},
                {"kind": "return", "from": 3, "to": 2, "label": "Historique trié"},
                {"kind": "return", "from": 2, "to": 1, "label": "Rows history"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/stats"},
                {"kind": "call", "from": 2, "to": 3, "label": "Calculer total / today / pending"},
                {"kind": "return", "from": 3, "to": 2, "label": "KPIs utilisateur"},
                {"kind": "return", "from": 2, "to": 1, "label": "Stats JSON"},
                {"kind": "return", "from": 1, "to": 0, "label": "Afficher cartes et activité récente"},
            ],
        },
        {
            "slug": "sequence_07_historique_documents",
            "title": "Diagramme de séquence - Consultation de l'historique documentaire",
            "description": "Récupération de l'historique depuis /api/documents/history.",
            "participants": ["Utilisateur", "DocumentHistory UI", "API Documents", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Accéder à /documents"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/history"},
                {"kind": "call", "from": 2, "to": 3, "label": "Filtrer par cin_number ou username"},
                {"kind": "return", "from": 3, "to": 2, "label": "Liste triée par created_at desc"},
                {"kind": "return", "from": 2, "to": 1, "label": "Rows normalisées"},
                {"kind": "return", "from": 1, "to": 0, "label": "Recherche, pagination, filtre"},
            ],
        },
        {
            "slug": "sequence_08_consulter_ma_cin",
            "title": "Diagramme de séquence - Consultation de la fiche Ma CIN",
            "description": "Scénario /api/documents/my-cin utilisé par l'écran MyCIN.",
            "participants": ["Utilisateur", "MyCIN UI", "API Documents", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /my-cin"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/my-cin"},
                {"kind": "call", "from": 2, "to": 3, "label": "Charger dernier DocumentRecord lié au CIN"},
                {"kind": "return", "from": 3, "to": 2, "label": "Document CIN ou vide"},
                {"kind": "return", "from": 2, "to": 1, "label": "Données d'identité normalisées"},
                {"kind": "return", "from": 1, "to": 0, "label": "Afficher carte d'identité"},
                {"kind": "note", "participant": 1, "label": "Cas alternatif: aucun document CIN -> message d'invitation à uploader la carte"},
            ],
        },
        {
            "slug": "sequence_09_upload_document_ocr",
            "title": "Diagramme de séquence - Upload document et pipeline OCR",
            "description": "Scénario /api/documents/upload avec OCR, extraction, classification et sauvegarde.",
            "participants": ["Utilisateur", "UploadDocument UI", "API Documents", "OCR Service", "Extractor Service", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Choisir fichier + uploader"},
                {"kind": "call", "from": 1, "to": 2, "label": "POST /documents/upload(formData)"},
                {"kind": "call", "from": 2, "to": 3, "label": "Envoyer le fichier à /ocr"},
                {"kind": "return", "from": 3, "to": 2, "label": "Texte OCR + lignes"},
                {"kind": "call", "from": 2, "to": 4, "label": "Envoyer text_lines à /extract"},
                {"kind": "return", "from": 4, "to": 2, "label": "Entités extraites"},
                {"kind": "call", "from": 2, "to": 2, "label": "Classifier + calculer quality_score"},
                {"kind": "call", "from": 2, "to": 5, "label": "saveToMongoDB(documentData)"},
                {"kind": "return", "from": 5, "to": 2, "label": "id sauvegardé"},
                {"kind": "return", "from": 2, "to": 1, "label": "success + document"},
                {"kind": "return", "from": 1, "to": 0, "label": "Résultat OCR affiché"},
                {"kind": "note", "participant": 2, "label": "Fallback prévu si le service OCR ou extractor n'est pas disponible"},
            ],
        },
        {
            "slug": "sequence_10_consulter_detail_document",
            "title": "Diagramme de séquence - Consultation du détail d'un document",
            "description": "Scénario /api/documents/:id avec contrôle d'accès utilisateur ou admin.",
            "participants": ["Utilisateur/Admin", "DocumentDetails UI", "API Documents", "Access Policy", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Choisir un document"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/:id"},
                {"kind": "call", "from": 2, "to": 4, "label": "findById(id)"},
                {"kind": "return", "from": 4, "to": 2, "label": "DocumentRecord"},
                {"kind": "call", "from": 2, "to": 3, "label": "canAccessDocument(user, document)"},
                {"kind": "return", "from": 3, "to": 2, "label": "allow / deny"},
                {"kind": "return", "from": 2, "to": 1, "label": "Document complet"},
                {"kind": "return", "from": 1, "to": 0, "label": "Affichage du détail"},
                {"kind": "note", "participant": 2, "label": "Cas alternatifs: document absent -> 404 | accès interdit -> 403"},
            ],
        },
        {
            "slug": "sequence_11_correction_entites",
            "title": "Diagramme de séquence - Correction des entités extraites",
            "description": "Scénario /api/documents/:id/entities depuis l'édition utilisateur ou admin.",
            "participants": ["Utilisateur/Admin", "Details UI", "API Documents", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Modifier entités / texte OCR"},
                {"kind": "call", "from": 1, "to": 2, "label": "PUT /documents/:id/entities"},
                {"kind": "call", "from": 2, "to": 3, "label": "Charger DocumentRecord"},
                {"kind": "return", "from": 3, "to": 2, "label": "Document existant"},
                {"kind": "call", "from": 2, "to": 2, "label": "Recalcul language + quality_score"},
                {"kind": "call", "from": 2, "to": 3, "label": "save() with edited_by_user=true"},
                {"kind": "return", "from": 3, "to": 2, "label": "Document mis à jour"},
                {"kind": "return", "from": 2, "to": 1, "label": "success + document"},
                {"kind": "return", "from": 1, "to": 0, "label": "Nouvelle version affichée"},
            ],
        },
        {
            "slug": "sequence_12_suppression_document",
            "title": "Diagramme de séquence - Suppression d'un document",
            "description": "Scénario /api/documents/:id avec suppression base + fichier disque.",
            "participants": ["Utilisateur/Admin", "DocumentHistory UI", "API Documents", "MongoDB", "FileSystem"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Cliquer Supprimer"},
                {"kind": "call", "from": 1, "to": 2, "label": "DELETE /documents/:id"},
                {"kind": "call", "from": 2, "to": 3, "label": "findById(id) + contrôle d'accès"},
                {"kind": "return", "from": 3, "to": 2, "label": "Document autorisé"},
                {"kind": "call", "from": 2, "to": 3, "label": "findByIdAndDelete(id)"},
                {"kind": "call", "from": 2, "to": 4, "label": "unlink(stored_filename) si présent"},
                {"kind": "return", "from": 2, "to": 1, "label": "success"},
                {"kind": "return", "from": 1, "to": 0, "label": "Historique rafraîchi"},
            ],
        },
        {
            "slug": "sequence_13_generation_notifications",
            "title": "Diagramme de séquence - Génération des notifications utilisateur",
            "description": "Notifications générées côté client à partir de l'historique et sauvegardées en localStorage.",
            "participants": ["Utilisateur", "Notifications UI", "API Documents", "LocalStorage"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /notifications"},
                {"kind": "call", "from": 1, "to": 3, "label": "Lire smartarchive.notifications"},
                {"kind": "return", "from": 3, "to": 1, "label": "État local ou vide"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /documents/history si aucun cache"},
                {"kind": "return", "from": 2, "to": 1, "label": "Rows history"},
                {"kind": "call", "from": 1, "to": 1, "label": "Mapper rows -> notifications"},
                {"kind": "call", "from": 1, "to": 3, "label": "Sauvegarder items JSON"},
                {"kind": "return", "from": 1, "to": 0, "label": "Liste filtrable affichée"},
            ],
        },
        {
            "slug": "sequence_14_admin_lister_utilisateurs",
            "title": "Diagramme de séquence - Admin: lister les utilisateurs",
            "description": "Scénario /api/admin/users pour l'écran UsersList / UsersManagement.",
            "participants": ["Administrateur", "Users UI", "API Admin", "User DB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /admin/users"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /admin/users"},
                {"kind": "call", "from": 2, "to": 3, "label": "find({}).sort(created_at desc)"},
                {"kind": "return", "from": 3, "to": 2, "label": "Liste users"},
                {"kind": "return", "from": 2, "to": 1, "label": "Rows normalisées"},
                {"kind": "return", "from": 1, "to": 0, "label": "Table utilisateurs"},
            ],
        },
        {
            "slug": "sequence_15_admin_modifier_utilisateur",
            "title": "Diagramme de séquence - Admin: modifier ou vérifier un utilisateur",
            "description": "Scénario /api/admin/users/:id avec mise à jour partielle des champs autorisés.",
            "participants": ["Administrateur", "UserDetails UI", "API Admin", "User DB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Modifier email / rôle / vérification"},
                {"kind": "call", "from": 1, "to": 2, "label": "PUT /admin/users/:id"},
                {"kind": "call", "from": 2, "to": 3, "label": "findByIdAndUpdate(id, updates)"},
                {"kind": "return", "from": 3, "to": 2, "label": "User mis à jour"},
                {"kind": "return", "from": 2, "to": 1, "label": "success + user"},
                {"kind": "return", "from": 1, "to": 0, "label": "Interface actualisée"},
            ],
        },
        {
            "slug": "sequence_16_admin_supprimer_utilisateur",
            "title": "Diagramme de séquence - Admin: suppression d'un utilisateur",
            "description": "Scénario /api/admin/users/:id en suppression définitive.",
            "participants": ["Administrateur", "Users UI", "API Admin", "User DB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Confirmer suppression"},
                {"kind": "call", "from": 1, "to": 2, "label": "DELETE /admin/users/:id"},
                {"kind": "call", "from": 2, "to": 3, "label": "findByIdAndDelete(id)"},
                {"kind": "return", "from": 3, "to": 2, "label": "deleted / not found"},
                {"kind": "return", "from": 2, "to": 1, "label": "success"},
                {"kind": "return", "from": 1, "to": 0, "label": "Liste rechargée"},
            ],
        },
        {
            "slug": "sequence_17_admin_revision_document",
            "title": "Diagramme de séquence - Admin: révision d'un document",
            "description": "Scénario /api/admin/documents/:id/review avec statut verified/rejected/pending_review/processed.",
            "participants": ["Administrateur", "DocumentDetails UI", "API Admin", "MongoDB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Choisir statut + note"},
                {"kind": "call", "from": 1, "to": 2, "label": "PUT /admin/documents/:id/review"},
                {"kind": "call", "from": 2, "to": 3, "label": "findByIdAndUpdate(status, review_note, reviewed_by)"},
                {"kind": "return", "from": 3, "to": 2, "label": "Document mis à jour"},
                {"kind": "return", "from": 2, "to": 1, "label": "success + document"},
                {"kind": "return", "from": 1, "to": 0, "label": "Résultat de validation affiché"},
            ],
        },
        {
            "slug": "sequence_18_admin_dashboard_stats",
            "title": "Diagramme de séquence - Admin: dashboard et statistiques",
            "description": "Chargement parallèle de /statistics, /users et /documents pour le dashboard admin.",
            "participants": ["Administrateur", "AdminDashboard UI", "API Admin", "User DB", "Document DB"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /admin/dashboard"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /admin/statistics"},
                {"kind": "call", "from": 2, "to": 3, "label": "count totalUsers / verifiedUsers"},
                {"kind": "call", "from": 2, "to": 4, "label": "count totalDocuments / documentsToday"},
                {"kind": "return", "from": 3, "to": 2, "label": "KPIs utilisateurs"},
                {"kind": "return", "from": 4, "to": 2, "label": "KPIs documents"},
                {"kind": "return", "from": 2, "to": 1, "label": "JSON statistiques"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /admin/users + GET /documents"},
                {"kind": "return", "from": 2, "to": 1, "label": "listes pour widgets"},
                {"kind": "return", "from": 1, "to": 0, "label": "Dashboard complet"},
            ],
        },
        {
            "slug": "sequence_19_admin_rapports_autocomplete",
            "title": "Diagramme de séquence - Admin: suggestions et autocomplétion de rapports",
            "description": "Scénario de ranking hybride Python/Ollama avec fallback JavaScript pour Reports.jsx.",
            "participants": ["Administrateur", "Reports UI", "API Admin", "User DB", "Document DB", "Python/Ollama"],
            "messages": [
                {"kind": "call", "from": 0, "to": 1, "label": "Saisir fragment de rapport"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /admin/reports/token-autocomplete"},
                {"kind": "call", "from": 2, "to": 3, "label": "Charger users pertinents"},
                {"kind": "call", "from": 2, "to": 4, "label": "Charger docs et entités"},
                {"kind": "return", "from": 3, "to": 2, "label": "Noms / CIN"},
                {"kind": "return", "from": 4, "to": 2, "label": "Certificats / champs OCR"},
                {"kind": "call", "from": 2, "to": 5, "label": "runPythonAutocomplete(...)"},
                {"kind": "return", "from": 5, "to": 2, "label": "Suggestions classées"},
                {"kind": "return", "from": 2, "to": 1, "label": "intent + suggestions"},
                {"kind": "return", "from": 1, "to": 0, "label": "UI d'aide à la rédaction"},
                {"kind": "note", "participant": 2, "label": "Fallback local JS si Python/Ollama indisponible"},
            ],
        },
        {
            "slug": "sequence_20_admin_parametres",
            "title": "Diagramme de séquence - Admin: consultation et mise à jour des paramètres",
            "description": "Scénarios /api/admin/settings GET puis PUT depuis SystemSettings / Settings.",
            "participants": ["Administrateur", "Settings UI", "API Admin", "AdminSettings DB"],
            "messages": [
                {"kind": "section", "label": "Lecture initiale"},
                {"kind": "call", "from": 0, "to": 1, "label": "Ouvrir /admin/settings"},
                {"kind": "call", "from": 1, "to": 2, "label": "GET /admin/settings"},
                {"kind": "call", "from": 2, "to": 3, "label": "findOne(key='global') ou create defaults"},
                {"kind": "return", "from": 3, "to": 2, "label": "Settings courants"},
                {"kind": "return", "from": 2, "to": 1, "label": "Payload de configuration"},
                {"kind": "section", "label": "Mise à jour"},
                {"kind": "call", "from": 0, "to": 1, "label": "Modifier timezone / thresholds / flags"},
                {"kind": "call", "from": 1, "to": 2, "label": "PUT /admin/settings"},
                {"kind": "call", "from": 2, "to": 3, "label": "save(updatedBy=admin)"},
                {"kind": "return", "from": 3, "to": 2, "label": "Settings mis à jour"},
                {"kind": "return", "from": 2, "to": 1, "label": "success + settings"},
            ],
        },
    ]


def render_use_case(spec):
    img, draw = make_canvas(spec["title"], spec["description"], size=(1800, 1100))
    boundary = (260, 130, 1560, 960)
    draw.rounded_rectangle(boundary, radius=28, outline=SLATE, width=3, fill=(251, 252, 254))
    draw_multiline(draw, (300, 140, 1520, 175), "Système SmartArchive", FONT_SMALL_BOLD, MUTED)

    registry = {}
    for actor in spec["actors"]:
        draw_actor(draw, actor["x"], actor["y"], actor["label"])
        registry[actor["id"]] = (actor["x"], actor["y"] + 80)

    usecase_rects = {}
    for usecase in spec["usecases"]:
        rect = usecase["rect"]
        draw_usecase(draw, rect, usecase["label"])
        usecase_rects[usecase["id"]] = rect
        registry[usecase["id"]] = center_of(rect)

    for source, target, kind in spec["links"]:
        start = registry[source]
        end = registry[target]
        label = ""
        dashed = False
        arrow = False
        color = LINE
        if kind == "include":
            label = "<<include>>"
            dashed = True
            arrow = True
            color = BLUE
        elif kind == "extend":
            label = "<<extend>>"
            dashed = True
            arrow = True
            color = VIOLET
        draw_connector(draw, start, end, label=label, color=color, dashed=dashed, arrow=arrow)

    target = ASSET_DIR / f"{spec['slug']}.png"
    img.save(target)
    return target


def render_class(spec):
    img, draw = make_canvas(spec["title"], spec["description"], size=(1760, 1260))
    rects = {}
    for class_spec in spec["classes"]:
        draw_class_box(
            draw,
            class_spec["rect"],
            class_spec["name"],
            class_spec["attrs"],
            class_spec["ops"],
            fill=class_spec["fill"],
            outline=class_spec["outline"],
        )
        rects[class_spec["id"]] = class_spec["rect"]

    for source, target, kind, label in spec["relations"]:
        source_rect = rects[source]
        target_rect = rects[target]
        start = anchor_point(source_rect, target_rect)
        end = anchor_point(target_rect, source_rect)
        color = BLUE if kind == "association" else GREEN if kind == "inheritance" else VIOLET
        draw_connector(
            draw,
            start,
            end,
            label=label,
            color=color,
            dashed=(kind == "dependency"),
            arrow=(kind in {"dependency", "inheritance"}),
            label_offset=-18,
        )

    target = ASSET_DIR / f"{spec['slug']}.png"
    img.save(target)
    return target


def render_sequence(spec):
    participant_count = len(spec["participants"])
    width = max(1800, 280 * participant_count + 200)
    base_height = 260
    for message in spec["messages"]:
        if message["kind"] == "section":
            base_height += 50
        elif message["kind"] == "note":
            base_height += 110
        else:
            base_height += 62
    height = max(1000, base_height + 120)

    img, draw = make_canvas(spec["title"], spec["description"], size=(width, height))
    x_positions = []
    left = 130
    right = width - 130
    if participant_count == 1:
        x_positions = [width // 2]
    else:
        step = (right - left) / (participant_count - 1)
        x_positions = [left + index * step for index in range(participant_count)]

    top = 130
    for index, participant in enumerate(spec["participants"]):
        draw_participant(draw, x_positions[index], top, participant)

    y = 240
    for message in spec["messages"]:
        kind = message["kind"]
        if kind == "section":
            draw.rounded_rectangle((80, y - 6, width - 80, y + 28), radius=10, fill=SLATE_FILL, outline=SLATE, width=2)
            draw_multiline(draw, (100, y - 4, width - 100, y + 26), message["label"], FONT_SMALL_BOLD, TEXT)
            y += 52
        elif kind in {"call", "return"}:
            x1 = x_positions[message["from"]]
            x2 = x_positions[message["to"]]
            draw_sequence_message(draw, x1, x2, y, message["label"], returned=(kind == "return"))
            y += 62
        elif kind == "note":
            participant_x = x_positions[message["participant"]]
            note_x = participant_x + 30
            if note_x + 380 > width - 30:
                note_x = participant_x - 410
            note_rect = draw_note(draw, int(note_x), y - 10, message["label"], width=380)
            y = note_rect[3] + 24

    target = ASSET_DIR / f"{spec['slug']}.png"
    img.save(target)
    return target


def add_cover(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Annexe\nDiagrammes complets SmartArchive")
    run.bold = True
    run.font.size = Pt(24)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "Cas d'utilisation, diagrammes de classes et diagrammes de séquence\n"
        "générés à partir des scénarios réels du projet."
    )
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("Projet analysé : Frontend-SmartArchive + Backend-SmartArchive")
    run.font.size = Pt(11)

    doc.add_page_break()


def setup_document(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    add_cover(doc)


def build_docx(manifest):
    doc = Document()
    setup_document(doc)

    doc.add_heading("Inventaire rapide", level=1)
    paragraph = doc.add_paragraph(style=None)
    paragraph.add_run(f"Total diagrammes: {len(manifest)}\n").bold = True
    use_case_count = len([item for item in manifest if item["category"] == "Cas d'utilisation"])
    class_count = len([item for item in manifest if item["category"] == "Classes"])
    sequence_count = len([item for item in manifest if item["category"] == "Séquences"])
    paragraph.add_run(f"- Cas d'utilisation: {use_case_count}\n")
    paragraph.add_run(f"- Classes: {class_count}\n")
    paragraph.add_run(f"- Séquences: {sequence_count}\n")
    paragraph.add_run(
        "Ces diagrammes couvrent les principaux scénarios codés: authentification, OCR CIN, "
        "upload documentaire, espace utilisateur, administration, reporting et paramètres."
    )
    doc.add_page_break()

    current_category = None
    for item in manifest:
        if item["category"] != current_category:
            current_category = item["category"]
            doc.add_heading(current_category, level=1)

        doc.add_heading(item["title"], level=2)
        doc.add_paragraph(item["description"])
        doc.add_picture(str(item["path"]), width=Cm(26))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(item["filename"])
        run.italic = True
        run.font.size = Pt(9)
        doc.add_page_break()

    output = choose_output_path(DOCX_OUTPUT)
    doc.save(str(output))
    return output


def write_index(manifest):
    lines = [
        "# SmartArchive - Index des diagrammes",
        "",
        f"Total: **{len(manifest)}** diagrammes.",
        "",
    ]
    current = None
    for item in manifest:
        if item["category"] != current:
            current = item["category"]
            lines.extend([f"## {current}", ""])
        lines.append(f"- `{item['filename']}` : {item['title']} - {item['description']}")
    INDEX_OUTPUT.write_text("\n".join(lines), encoding="utf-8")


def build_manifest():
    ensure_dir()
    manifest = []

    for spec in use_case_specs():
        path = render_use_case(spec)
        manifest.append(
            {
                "category": "Cas d'utilisation",
                "title": spec["title"],
                "description": spec["description"],
                "filename": path.name,
                "path": path,
            }
        )

    for spec in class_specs():
        path = render_class(spec)
        manifest.append(
            {
                "category": "Classes",
                "title": spec["title"],
                "description": spec["description"],
                "filename": path.name,
                "path": path,
            }
        )

    for spec in sequence_specs():
        path = render_sequence(spec)
        manifest.append(
            {
                "category": "Séquences",
                "title": spec["title"],
                "description": spec["description"],
                "filename": path.name,
                "path": path,
            }
        )

    serializable = [
        {
            key: (str(value) if isinstance(value, Path) else value)
            for key, value in item.items()
        }
        for item in manifest
    ]
    MANIFEST_OUTPUT.write_text(
        json.dumps(serializable, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_index(manifest)
    return manifest


def main():
    manifest = build_manifest()
    output = build_docx(manifest)
    print(f"Diagrammes générés : {ASSET_DIR}")
    print(f"Annexe générée : {output}")
    print(f"Total diagrammes : {len(manifest)}")


if __name__ == "__main__":
    main()
